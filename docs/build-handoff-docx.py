#!/usr/bin/env python3
"""
Genera docs/traspaso-conocimiento.docx — documento de traspaso del proyecto
asisa-pc para el equipo que va a migrarlo de Vercel a Azure.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'traspaso-conocimiento.docx')

doc = Document()

# --- Estilos base ---
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

# Estilo "Code" (monospace) — creamos un Character Style nuevo
try:
    code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(9)
    code_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    code_style.paragraph_format.space_before = Pt(2)
    code_style.paragraph_format.space_after = Pt(2)
except Exception:
    code_style = doc.styles['Normal']


def add_code_block(text):
    for line in text.split('\n'):
        p = doc.add_paragraph(line if line else ' ', style='Code')
        # Fondo gris claro
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'F4F4F4')
        shd.set(qn('w:val'), 'clear')
        pPr.append(shd)


def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def add_bullet(text):
    return doc.add_paragraph(text, style='List Bullet')


def add_numbered(text):
    return doc.add_paragraph(text, style='List Number')


def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    if widths:
        for i, w in enumerate(widths):
            for cell in t.columns[i].cells:
                cell.width = w


# =============================================================================
# Portada
# =============================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('Proyecto asisa-pc')
title_run.font.size = Pt(28)
title_run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('Traspaso de conocimiento y plan de migración a Azure')
subtitle_run.font.size = Pt(16)
subtitle_run.italic = True

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Equipo: Softtek · ASISA\n').bold = True
meta.add_run('Versión: 1.0 · Fecha: 2026-05-27\n')
meta.add_run('Stack actual: AEM Author + Edge Delivery Services + Vercel\n')
meta.add_run('Stack objetivo: AEM Author + Edge Delivery Services + Azure')

doc.add_page_break()

# =============================================================================
# 0. Resumen ejecutivo
# =============================================================================
doc.add_heading('0. Resumen ejecutivo', level=1)
doc.add_paragraph(
    'asisa-pc es el sitio del Cuadro Médico de ASISA construido sobre Adobe Edge Delivery '
    'Services (EDS). Sirve URLs ultrarrápidas indexables por Google a partir de una arquitectura '
    'híbrida que combina:'
)
add_bullet('AEM Author (Adobe Cloud) — autoría de plantillas y contenido estático.')
add_bullet('Edge Delivery Services (aem.live / aem.page) — CDN y capa de entrega de Adobe.')
add_bullet('Vercel — capa "BYOM" (Bring Your Own Markup): genera bajo demanda el HTML de '
           'decenas de miles de URLs dinámicas (provincias, especialidades, fichas de médico/centro) '
           'y expone unas APIs JSON que consumen los bloques en cliente.')
add_bullet('GitHub (asisa-softtek/asisa-pc) — fuente de verdad del código y de los JSON cacheados '
           'que alimentan la web.')

doc.add_paragraph(
    'El objetivo de este documento es doble: (1) traspasar el conocimiento completo del proyecto '
    'al equipo entrante y (2) servir como guía para sustituir Vercel por Azure manteniendo intacto '
    'el resto del stack.'
)
doc.add_paragraph(
    'Volumen aproximado a 2026-05: ~36.000 URLs dinámicas en el cuadro médico (52 provincias, '
    '~3.250 combinaciones provincia+especialidad, ~20.500 doctores, ~6.500 centros, ~181 '
    'especialidades nacionales).'
)

# =============================================================================
# 1. Arquitectura general
# =============================================================================
doc.add_heading('1. Arquitectura general', level=1)
doc.add_paragraph('Flujo de una petición de usuario a una URL dinámica '
                  '(ej. /cuadro-medico/p/madrid/pe/cardiologia):')
add_code_block(
    'Usuario\n'
    '   │  GET https://main--asisa-pc--asisa-softtek.aem.live/cuadro-medico/p/madrid/pe/cardiologia\n'
    '   ▼\n'
    'CDN de Edge Delivery Services\n'
    '   │  Sirve HTML ya procesado (si está en el bus live).\n'
    '   ▼\n'
    'Navegador ejecuta scripts/aem.js + scripts/scripts.js\n'
    '   │  Detecta bloques <div class="cuadro-medico"> y los hidrata.\n'
    '   ▼\n'
    'Bloque JS lee window.location.pathname → llama a la API de datos\n'
    '   │  fetch("https://asisa-pc.vercel.app/api/providers?provinceSlug=madrid&specSlug=cardiologia&...")\n'
    '   ▼\n'
    'Función serverless en Vercel\n'
    '   │  Lee JSON cacheados de data/providers/madrid/cardiologia.json.\n'
    '   ▼\n'
    'Devuelve el listado paginado al navegador, que renderiza tarjetas.'
)
doc.add_paragraph('¿Y cómo entró ese HTML en la CDN de EDS la primera vez?')
add_numbered('Alguien hizo POST /preview/.../cuadro-medico/p/madrid/pe/cardiologia.')
add_numbered('EDS, al ver que es una URL dinámica, consultó al overlay en Vercel: '
             'GET https://asisa-pc.vercel.app/markup/cuadro-medico/p/madrid/pe/cardiologia.')
add_numbered('Vercel respondió con una plantilla HTML mínima que tiene los <div class="..."> '
             'de los bloques (cuadro-medico, otras-especialidades, otras-provincias) vacíos.')
add_numbered('EDS guardó ese HTML en su bus preview, lo procesó (extracción de secciones y '
             'bloques) y lo dejó listo para servir.')
add_numbered('Un POST /live posterior promocionó el HTML de preview al bus live para que sea '
             'público.')

# =============================================================================
# 2. Setup inicial EDS — los comandos críticos
# =============================================================================
doc.add_heading('2. Setup inicial EDS (configuración del Config Service)', level=1)
doc.add_paragraph(
    'EDS se configura mediante POSTs al "Config Service" en admin.hlx.page. '
    'Esto se hace UNA VEZ por entorno y solo se repite si cambia la arquitectura. '
    'Para modificarlos hace falta un token con rol config_admin (cuenta técnica de Adobe; '
    'jorge.lorenzo@ext.softtek.com sólo tiene rol admin y NO puede tocar estos endpoints).'
)
doc.add_paragraph(
    'IMPORTANTE — Tokens en este proyecto: el access.json (§2.3) tiene "requireAuth": "auto", '
    'lo que significa que las operaciones de preview, live, code, index y status admiten '
    'llamadas anónimas. Por tanto, todos los refrescos del día a día (POST /preview, /live, '
    '/code, /index; DELETE; GET /status) funcionan SIN token. Solo necesitan token los '
    'endpoints de configuración (POST /config/...) y los scripts que hablan directamente con '
    'AEM Author (create-aem-pages.mjs requiere AEM_TOKEN).'
)

doc.add_heading('2.1 Sitewide config (vincula GitHub, AEM y el overlay)', level=2)
add_code_block(
    "curl --request POST \\\n"
    "  --url https://admin.hlx.page/config/asisa-softtek/sites/asisa-pc.json \\\n"
    "  --header 'Content-Type: application/json' \\\n"
    "  --header 'x-auth-token: <TOKEN_CONFIG_ADMIN>' \\\n"
    "  --data '{\n"
    '    "code": {\n'
    '      "owner": "asisa-softtek",\n'
    '      "repo": "asisa-pc",\n'
    '      "source": { "type": "github", "url": "https://github.com/asisa-softtek/asisa-pc" }\n'
    '    },\n'
    '    "content": {\n'
    '      "source": {\n'
    '        "url": "https://author-p133185-e1320482.adobeaemcloud.com/bin/franklin.delivery/asisa-softtek/asisa-pc/main",\n'
    '        "type": "markup",\n'
    '        "suffix": ".html"\n'
    '      },\n'
    '      "overlay": {\n'
    '        "url": "https://asisa-pc.vercel.app/markup",\n'
    '        "type": "markup"\n'
    '      }\n'
    "    }\n"
    "  }'"
)
doc.add_paragraph(
    'Qué hace: EDS, para cada request, consulta primero al overlay. Si responde 200, usa esa '
    'respuesta. Si responde 404, hace fallback al source (AEM Author). En la migración a Azure, '
    'lo único que hay que cambiar aquí es la URL del overlay.'
)

doc.add_heading('2.2 Public config (mappings de rutas dinámicas)', level=2)
add_code_block(
    "curl --request POST \\\n"
    "  --url https://admin.hlx.page/config/asisa-softtek/sites/asisa-pc/public.json \\\n"
    "  --header 'Content-Type: application/json' \\\n"
    "  --header 'x-auth-token: <TOKEN_CONFIG_ADMIN>' \\\n"
    "  --data '{\n"
    '    "paths": {\n'
    '      "mappings": [\n'
    '        "/content/site-pc/:/",\n'
    '        "/content/site-pc/configuration:/.helix/config.json",\n'
    '        "/cuadro-medico/p/*:/cuadro-medico/provincia",\n'
    '        "/cuadro-medico/d/*:/cuadro-medico/doctor",\n'
    '        "/cuadro-medico/e/*:/cuadro-medico/especialidad",\n'
    '        "/cuadro-medico/c/*:/cuadro-medico/centro"\n'
    '      ],\n'
    '      "includes": ["/content/site-pc/"]\n'
    "    }\n"
    "  }'"
)
doc.add_paragraph(
    'Define la red de seguridad para fallback a AEM: si el overlay devuelve 404 para una URL '
    'dinámica, EDS sabe qué plantilla cargar desde AEM Author.'
)

doc.add_heading('2.3 Access config (roles admin / config_admin)', level=2)
add_code_block(
    "curl --request POST \\\n"
    "  --url https://admin.hlx.page/config/asisa-softtek/sites/asisa-pc/access.json \\\n"
    "  --header 'Content-Type: application/json' \\\n"
    "  --header 'x-auth-token: <TOKEN_CONFIG_ADMIN>' \\\n"
    "  --data '{\n"
    '    "admin": {\n'
    '      "role": {\n'
    '        "admin": ["jorge.lorenzo@ext.softtek.com"],\n'
    '        "config_admin": [\n'
    '          "662F1E56661D006D0A495E33@techacct.adobe.com",\n'
    '          "jorge.lorenzo@ext.softtek.com"\n'
    '        ]\n'
    '      },\n'
    '      "requireAuth": "auto"\n'
    "    }\n"
    "  }'"
)
doc.add_paragraph(
    'admin permite previewar, publicar y refrescar code. config_admin permite ADEMÁS modificar '
    'estos 3 endpoints. Sin config_admin, todos los POST /config/... devuelven 401.'
)

doc.add_heading('2.4 Headers config (opcional pero recomendado en migración Azure)', level=2)
doc.add_paragraph(
    'Si Azure necesita un header específico para autenticar al overlay (API key, bypass token, '
    'etc.), se inyecta vía headers.json:'
)
add_code_block(
    "curl --request POST \\\n"
    "  --url https://admin.hlx.page/config/asisa-softtek/sites/asisa-pc/headers.json \\\n"
    "  --header 'Content-Type: application/json' \\\n"
    "  --header 'x-auth-token: <TOKEN_CONFIG_ADMIN>' \\\n"
    "  --data '{\n"
    '    "mappings": [\n'
    '      {\n'
    '        "path": "/**",\n'
    '        "headers": { "X-BYOM-Origin": "Azure-Functions" }\n'
    "      }\n"
    "    ]\n"
    "  }'"
)

doc.add_heading('2.5 Verificación', level=2)
add_code_block(
    "# Anónimo (no necesita token gracias a requireAuth: auto)\n"
    "curl https://admin.hlx.page/status/asisa-softtek/asisa-pc/main/cuadro-medico/p/madrid"
)
doc.add_paragraph(
    'Debe responder con un sourceLocation que apunte al overlay '
    '("markup:https://<overlay-host>/markup/cuadro-medico/p/madrid"). Si apunta al author de AEM, '
    'el overlay no está aplicado.'
)

doc.add_heading('2.6 Resumen de tokens en el proyecto', level=2)
add_table(
    ['Operación', '¿Token?', 'Cuál'],
    [
        ('POST /config/asisa-softtek/sites/*.json', 'Sí', 'x-auth-token con rol config_admin '
         '(cuenta técnica de Adobe)'),
        ('POST /preview, /live, /code, /index', 'No', 'Anónimo (requireAuth: auto)'),
        ('DELETE /preview, /live', 'No', 'Anónimo'),
        ('GET /status, /profile', 'No', 'Anónimo'),
        ('POST /bin/wcmcommand, /bin/replicate.json (AEM Author)', 'Sí',
         'AEM_TOKEN (cookie login-token del Author)'),
        ('GET /api/sync-aem (endpoint propio en Vercel)', 'Sí',
         'SYNC_SECRET como query param'),
        ('Cualquier otro /api/* (Vercel)', 'No', 'Públicos con CORS abierto'),
    ],
    widths=[Inches(2.5), Inches(0.8), Inches(3.2)],
)

# =============================================================================
# 3. Estructura del repositorio
# =============================================================================
doc.add_heading('3. Estructura del repositorio (file-by-file)', level=1)
doc.add_paragraph(
    'El repo asisa-softtek/asisa-pc contiene tanto el código del frontend EDS como las funciones '
    'serverless del backend Vercel (carpeta api/). En la migración, todo el código del frontend '
    '(blocks/, scripts/, styles/, head.html, configuraciones EDS) se queda igual; solo se reemplaza '
    'la carpeta api/ y los ficheros de configuración de Vercel.'
)

doc.add_heading('3.1 Configuración EDS / AEM', level=2)
add_table(
    ['Fichero', 'Función'],
    [
        ('fstab.yaml', 'Mountpoints. Solo monta "/" desde AEM Author '
                       '(/bin/franklin.delivery/asisa-softtek/asisa-pc/main, sufijo .html). Los '
                       'sitemaps ya NO se montan aquí — los genera EDS nativamente vía '
                       'helix-sitemap.yaml.'),
        ('helix-query.yaml', '6 índices de pages: pages (estáticas), cuadro-medico-provincias, '
                             'cuadro-medico-provincia-specs, cuadro-medico-doctores, '
                             'cuadro-medico-centros, cuadro-medico-especialidades. Cada uno '
                             'produce un /query-index-*.json filtrado por glob patterns.'),
        ('helix-sitemap.yaml', '6 sitemaps, uno por índice, con origin: https://www.asisa.es '
                               'para que las URLs en el XML apunten a producción.'),
        ('paths.json', 'Mapping interno: /content/site-pc/ → /.'),
        ('head.html', 'HTML inyectado en TODAS las páginas. Contiene CSP, scripts principales '
                      '(aem.js, scripts.js), styles.css local y las clientlibs CSS de ASISA '
                      'servidas por proxy vía Vercel /etc.clientlibs/*.'),
        ('404.html', 'Página de error custom con botón "volver" y evento RUM.'),
        ('.hlxignore', 'Excluye dotfiles y similares de la publicación EDS.'),
        ('tools/sidekick/config.json', 'Config del Sidekick (toolbar de edición).'),
    ],
    widths=[Inches(2), Inches(4.5)],
)

doc.add_heading('3.2 Configuración Vercel (la que cambia en la migración)', level=2)
add_table(
    ['Fichero', 'Función'],
    [
        ('vercel.json', 'Headers y rewrites. CORS abierto en /etc.clientlibs/*; rewrites para '
                        '/markup/*, /sitemap*.xml, /api/markup → /api/markup. Proxy '
                        '/etc.clientlibs/* → www.asisa.es/etc.clientlibs/* (clientlibs del DS '
                        'corporativo).'),
        ('.vercel/project.json', 'IDs internos del proyecto Vercel.'),
    ],
    widths=[Inches(2), Inches(4.5)],
)

doc.add_heading('3.3 Universal Editor — modelos de bloques (AEM Author)', level=2)
add_table(
    ['Fichero', 'Función'],
    [
        ('component-definition.json', 'Catálogo de componentes (Default Content, Sections, '
                                       'Blocks incluyendo los 8 bloques de cuadro médico).'),
        ('component-filters.json', '¿Qué bloque puede ir dentro de qué contenedor?'),
        ('component-models.json', 'Esquema de campos editables por componente.'),
        ('models/_*.json', 'Fragmentos modulares de las definiciones globales.'),
    ],
    widths=[Inches(2), Inches(4.5)],
)

doc.add_heading('3.4 Endpoints API (carpeta api/) — funciones serverless en Vercel', level=2)
doc.add_paragraph(
    'Todos los ficheros api/*.js son funciones serverless Node.js (ESM puro, sin dependencias en '
    'runtime). En Vercel, cada uno se expone como una function en /api/<nombre>. Algunos están '
    'remapeados por vercel.json (ej. /markup/* → /api/markup).'
)
add_table(
    ['Fichero', 'Ruta HTTP', 'Función'],
    [
        ('api/markup.js', '/markup/* (rewrite)',
         'Overlay BYOM. Devuelve plantillas HTML para /cuadro-medico/p/*, /d/*, /c/*, /e/*. 404 '
         'para el resto. También sirve los 6 sitemaps en /markup/sitemap*.xml.'),
        ('api/providers.js', '/api/providers',
         'Listado paginado de profesionales/centros (provinceSlug, specSlug, tab, page, limit).'),
        ('api/doctor.js', '/api/doctor?key=…',
         'Ficha completa de profesional.'),
        ('api/centro.js', '/api/centro?key=…',
         'Ficha completa de centro.'),
        ('api/provincias.js', '/api/provincias[?slug=]',
         'Listado de provincias o detalle de una.'),
        ('api/especialidades.js', '/api/especialidades[?slug=]',
         'Master list de especialidades o detalle.'),
        ('api/providers-detail.js', '/api/providers-detail?id=…',
         'Detalle consolidado por código de localización.'),
        ('api/specialities.js', '/api/specialities?provinceCode=…',
         'PROXY en vivo al backend ASISA (ursaepre.asisa.es) — no usa caché local.'),
        ('api/sitemap.js', '/sitemap.xml (rewrite)',
         'Sitemap index "Vercel-only" (no usado por EDS). Apunta a www.asisa.es.'),
        ('api/sitemap-cuadro-medico.js', '/sitemap-cuadro-medico-*.xml (rewrite)',
         'Sitemap específico por tipo. EDS no lo usa, pero está accesible directo en Vercel.'),
        ('api/sync-aem.js', '/api/sync-aem',
         'Sincroniza URLs del sitemap remoto de www.asisa.es con preview+live de EDS. Requiere '
         'SYNC_SECRET y HLX_ADMIN_API_TOKEN.'),
    ],
    widths=[Inches(1.8), Inches(1.8), Inches(2.9)],
)
doc.add_paragraph(
    'Importante: varios endpoints implementan caché in-memory (variables module-scope *Cache) '
    'para evitar releer los JSON en cada petición. En Azure Functions hay que tener en cuenta que '
    'estas cachés solo persisten dentro de la misma instancia warm; en cold start se reconstruyen.'
)

doc.add_page_break()
doc.add_heading('3.4.1 Detalle exhaustivo de cada endpoint API', level=3)
doc.add_paragraph(
    'A continuación se documenta cada uno de los 11 endpoints de la carpeta api/: fichero, '
    'método, query params con tipo y validaciones, ficheros de data/ que lee, cachés in-memory, '
    'headers de respuesta, shape EXACTA del JSON o XML devuelto, y reglas de negocio. Esta es la '
    'información imprescindible para portar cada función serverless a Azure Functions sin perder '
    'comportamiento.'
)

# --- api/markup.js ---
doc.add_heading('Endpoint 1 · api/markup.js (overlay BYOM)', level=4)
doc.add_paragraph('Ruta: api/markup.js · 135 líneas')
doc.add_paragraph(
    'HTTP: GET. Vercel reescribe /markup y /markup/(.*) → /api/markup?path=$1. '
    'Devuelve HTML (text/html) para URLs dinámicas o XML (text/xml) para sitemaps. 404 con '
    'header "x-source: overlay-pass" si el path no coincide con ninguna regla — eso indica a EDS '
    'que pase al source de AEM.'
)
doc.add_paragraph('Query params:')
add_table(
    ['Param', 'Tipo', 'Obligatorio', 'Descripción'],
    [('path', 'String', 'Sí',
      'Ruta del recurso solicitado (sin el prefijo /markup). Se acepta con o sin sufijo .html o '
      '.plain.html — el handler los normaliza.')],
    widths=[Inches(0.8), Inches(0.8), Inches(0.8), Inches(4.1)],
)
doc.add_paragraph('Patrones de path reconocidos:')
add_bullet('/cuadro-medico/p/<slug> → template provincia (3 bloques)')
add_bullet('/cuadro-medico/p/<prov>/pe/<spec> → mismo template (los bloques leen el path)')
add_bullet('/cuadro-medico/d/<key> → template doctor (2 bloques)')
add_bullet('/cuadro-medico/c/<key> → template centro (1 bloque)')
add_bullet('/cuadro-medico/e/<slug> → template especialidad (3 bloques)')
add_bullet('/sitemap.xml → sitemap index (delega en api/sitemap.js)')
add_bullet('/sitemap-cuadro-medico-{provincias|provincia-specs|doctores|centros|especialidades}.xml '
           '→ delega en api/sitemap-cuadro-medico.js')
add_bullet('Cualquier otro → 404 overlay-pass')
doc.add_paragraph('Headers de respuesta:')
add_bullet('Content-Type: text/html; charset=utf-8 (templates) o application/xml; charset=utf-8 (sitemaps)')
add_bullet('Cache-Control: public, max-age=60 (EDS lo cachea aparte; este TTL es para Vercel CDN)')
add_bullet('x-source: template:provincia | template:doctor | template:centro | '
           'template:especialidad | sitemap:index | sitemap:<type> | overlay-pass')
doc.add_paragraph(
    'No lee data/ directamente. Importa las funciones getSitemapIndexXml() de api/sitemap.js y '
    'getCuadroMedicoSitemapXml(type) de api/sitemap-cuadro-medico.js para servir los sitemaps. '
    'Helpers: extractPath() (normaliza el query param), isProvinciaPath/isDoctorPath/'
    'isCentroPath/isEspecialidadPath (regex matchers).'
)
doc.add_paragraph(
    'Inyecta un <script> con history.replaceState() en el HTML para que, al abrir el overlay '
    'directamente desde Vercel (/markup/cuadro-medico/p/madrid), los bloques JS vean el path real '
    'cuando leen window.location.pathname.'
)
doc.add_paragraph('Consumido por: EDS preview/live al hacer fetch al overlay configurado en sitewide config.')

# --- api/providers.js ---
doc.add_heading('Endpoint 2 · api/providers.js (listado paginado)', level=4)
doc.add_paragraph('Ruta: api/providers.js · 194 líneas · HTTP: GET /api/providers')
add_table(
    ['Param', 'Tipo', 'Obligatorio', 'Default', 'Descripción'],
    [
        ('provinceSlug', 'String', 'Opcional¹', '-', 'Slug de provincia.'),
        ('specSlug', 'String', 'Opcional¹', '-', 'Slug de especialidad.'),
        ('tab', "'professionals'|'centers'", 'Opcional', 'professionals', 'Pestaña.'),
        ('page', 'Number', 'Opcional', '1', 'Número de página (1-indexed).'),
        ('limit', 'Number', 'Opcional', '10', 'Tamaño de página. Capado a MAX_LIMIT=50.'),
    ],
    widths=[Inches(1.2), Inches(1.5), Inches(0.9), Inches(0.8), Inches(2.1)],
)
doc.add_paragraph(
    '¹ Al menos UNO de provinceSlug o specSlug es obligatorio. Si se pasa solo specSlug, hace '
    'búsqueda nacional escaneando todas las provincias.'
)
doc.add_paragraph('Errores HTTP:')
add_bullet('400 { error: "provinceSlug or specSlug is required" }')
add_bullet('400 { error: "tab must be professionals or centers" }')
add_bullet('404 { error: "Province not found: <slug>" }')
add_bullet('404 { error: "Speciality not found: <slug>" }')
add_bullet('404 { error: "No data for <prov>/<spec>" }')
doc.add_paragraph('Ficheros leídos: data/provincias.json, data/cuadro-medico/especialidades.json, '
                  'data/providers/<provinceSlug>/<specSlug>.json y/o todos los ficheros de '
                  'data/providers/<provinceSlug>/*.json (para modo "solo provincia").')
doc.add_paragraph('Cachés in-memory: provinciasCache (master), especialidadesCache (master), '
                  'allProvinceCache (Map<provSlug, lista_deduplicada>).')
doc.add_paragraph('Headers: Access-Control-Allow-Origin: *, Cache-Control: public, '
                  's-maxage=300, stale-while-revalidate=3600 (5 min CDN + 1h stale).')
doc.add_paragraph('Shape de la respuesta (200):')
add_code_block(
    '{\n'
    '  "provinceSlug": "madrid" | null,\n'
    '  "specSlug": "cardiologia" | null,\n'
    '  "tab": "professionals" | "centers",\n'
    '  "page": 1,\n'
    '  "limit": 10,\n'
    '  "total": Number,                 // total CAPADO (max 30 prof, 50 centros)\n'
    '  "totalPages": Number,\n'
    '  "totalProfessionals": Number,    // count CAPADO a 30\n'
    '  "totalCenters": Number,          // count CAPADO a 50\n'
    '  "results": [\n'
    '    {\n'
    '      "name": "DR. PEREZ GARCIA",\n'
    '      "speciality": "CARDIOLOGÍA",\n'
    '      "providerType": Number|null,\n'
    '      "doctorType": Number|null,   // 1 = profesional\n'
    '      "businessGroup": Boolean,\n'
    '      "parentDescription": "Hospital HM …",\n'
    '      "address": "DOCTOR ESQUERDO 83",\n'
    '      "postalCode": "28028",\n'
    '      "city": "MADRID",\n'
    '      "phone": "915732022",\n'
    '      "lat": 40.43,\n'
    '      "lon": -3.66,\n'
    '      "onlineAppointment": Boolean,\n'
    '      "videoConsultation": Boolean,\n'
    '      "ePrescription": Boolean,\n'
    '      "languages": ["es", "en"],\n'
    '      "collegiateCode": "281234567",\n'
    '      "providerLocalicationCode": 1234567,\n'
    '      "providerCode": 7654321,\n'
    '      "detailUrl": "/cuadro-medico/d/<slug>-<id>"\n'
    '    }\n'
    '  ]\n'
    '}'
)
doc.add_paragraph('Lógica de negocio:')
add_bullet('Filtra por tab: isProfessional() = doctorType === "1"; '
           'isCenter() = providerType ∈ {3,4,8,2,9}.')
add_bullet('Aplica caps MAX_TOTAL_BY_TAB: 30 prof / 50 centros antes de paginar.')
add_bullet('Modo "solo provincia" (sin specSlug): lee TODOS los ficheros de '
           'data/providers/<prov>/, deduplica por providerCode|providerLocalicationCode.')
add_bullet('Modo "nacional" (sin provinceSlug): escanea todas las carpetas data/providers/*/ '
           'buscando el spec.')
add_bullet('Consumido por bloques: cuadro-medico (motor búsqueda) y cuadro-medico-otros-medicos '
           '(con limit=50).')

# --- api/doctor.js ---
doc.add_heading('Endpoint 3 · api/doctor.js (ficha de profesional)', level=4)
doc.add_paragraph('Ruta: api/doctor.js · 161 líneas · HTTP: GET /api/doctor')
add_table(
    ['Param', 'Tipo', 'Obligatorio', 'Descripción'],
    [('key', 'String', 'Sí',
      'Identificador único del doctor en doctores-index.json '
      '(formato: "<slug-nombre>-<collegiateCode|providerCode>").')],
    widths=[Inches(0.8), Inches(0.8), Inches(0.8), Inches(4.1)],
)
doc.add_paragraph('Errores: 400 si falta key; 404 "Doctor not found"; 404 "No locations for".')
doc.add_paragraph('Ficheros leídos: doctores-index.json, provider-details/<locCode>.json (por cada '
                  'ubicación), providers/<prov>/<spec>.json (cross-ref).')
doc.add_paragraph('Cachés: indexCache, providersListCache (Map<"prov|spec", lista>).')
doc.add_paragraph('Cache-Control: public, s-maxage=3600, stale-while-revalidate=86400.')
doc.add_paragraph('Shape (resumen):')
add_code_block(
    '{\n'
    '  "key": "perez-garcia-juan-1234567",\n'
    '  "name": "PEREZ GARCIA, JUAN",\n'
    '  "collegiateCode": "281234567",\n'
    '  "languages": ["es", "en"],\n'
    '  "specialities": ["CARDIOLOGÍA", "MEDICINA INTERNA"],\n'
    '  // -- Campos del "representative" (primer location con detail) --\n'
    '  "specSlug", "provinceSlug", "parentDescription", "address",\n'
    '  "postalCode", "city", "provinceCode", "phone", "lat", "lon",\n'
    '  "onlineAppointment", "videoConsultation", "ePrescription",\n'
    '  "businessGroup",\n'
    '  "tuotempo": { presential, online, video, phone, asisaLive },\n'
    '  // -- Todas las ubicaciones donde trabaja --\n'
    '  "locations": [\n'
    '    { providerCode, providerLocalicationCode, specSlug, provinceSlug,\n'
    '      speciality, parentDescription, businessGroup, address, postalCode,\n'
    '      city, provinceCode, phone, lat, lon, onlineAppointment,\n'
    '      videoConsultation, ePrescription, tuotempo }\n'
    '  ]\n'
    '}'
)
doc.add_paragraph('Lógica clave:')
add_bullet('pickRepresentative(): elige la ubicación que tenga fichero de detalle si existe.')
add_bullet('mergeAddress(): combina address de detail y de la lista, priorizando valores '
           '"meaningful" (no vacíos, no 0).')
add_bullet('Una persona puede tener varias entradas (cada ubicación + especialidad genera una); '
           'se agrupa por collegiateCode en el índice.')
add_bullet('Consumido por: cuadro-medico-ficha-doctor, cuadro-medico-otros-medicos, '
           'cuadro-medico-spec-localizacion (ya retirado).')

# --- api/centro.js ---
doc.add_heading('Endpoint 4 · api/centro.js (ficha de centro)', level=4)
doc.add_paragraph('Ruta: api/centro.js · 259 líneas · HTTP: GET /api/centro')
doc.add_paragraph('Query: key (String, obligatorio) — slug del centro en centros-index.json.')
doc.add_paragraph('Errores: 400 si falta key; 404 "Centro not found"; 404 "Centro has no data".')
doc.add_paragraph('Ficheros leídos: centros-index.json, doctores-index.json (para validar qué '
                  'médicos tienen ficha publicada), data/providers/<prov>/*.json (scan completo).')
doc.add_paragraph('Cachés: centrosIndexCache, doctoresIndexCache, provinceScanCache '
                  '(Map<provSlug, { centros, doctorsByParent }>).')
doc.add_paragraph('Shape (resumen):')
add_code_block(
    '{\n'
    '  "key", "providerLocalicationCode", "name", "providerType",\n'
    '  "businessGroup", "address", "postalCode", "city", "provinceCode",\n'
    '  "provinceSlug", "phone", "lat", "lon",\n'
    '  "onlineAppointment", "videoConsultation", "ePrescription",\n'
    '  "specialities": [\n'
    '    {\n'
    '      "speciality": "CARDIOLOGÍA",\n'
    '      "specSlug": "cardiologia",\n'
    '      "phone", "onlineAppointment", "videoConsultation", "ePrescription",\n'
    '      "subSpecialities": ["..."],\n'
    '      "doctors": [{ key, name, subSpeciality }],\n'
    '      "observations": ""\n'
    '    }\n'
    '  ],\n'
    '  "doctors": [{ key, name, speciality, gender }],\n'
    '  "otherCentros": [\n'
    '    { key, providerLocalicationCode, name, providerType, businessGroup,\n'
    '      address, postalCode, city, provinceCode, phone, lat, lon,\n'
    '      specialities: ["..."],   // máx 4\n'
    '      specialitiesMore: Number  // el resto (n - 4)\n'
    '    }\n'
    '  ],\n'
    '  "description": "Centro con N especialidades en {ciudad}…"\n'
    '}'
)
doc.add_paragraph('Lógica clave:')
add_bullet('scanProvince(provSlug): parser que recorre TODOS los .json de '
           'data/providers/<prov>/, agrupa por centro y por doctor-padre. Cacheado.')
add_bullet('Vincula doctors al centro vía parentCode → providerLocalicationCode.')
add_bullet('otherCentros: top 4 centros con MAYOR overlap de especialidades (orden: overlap '
           'desc, nombre asc).')
add_bullet('Consumido por: cuadro-medico-ficha-centro.')

# --- api/provincias.js ---
doc.add_heading('Endpoint 5 · api/provincias.js (listado / detalle de provincia)', level=4)
doc.add_paragraph('Ruta: api/provincias.js · 20 líneas · HTTP: GET /api/provincias')
doc.add_paragraph('Query: slug (opcional). Sin slug → array de provincias. Con slug → detalle.')
doc.add_paragraph('Sin caché in-memory (passthrough de fichero). Cache-Control: s-maxage=86400, '
                  'stale-while-revalidate=604800 (24h + 7d stale).')
doc.add_paragraph('Shape sin slug:')
add_code_block(
    '[\n'
    '  { "name": "MADRID", "displayName": "Madrid", "slug": "madrid",\n'
    '    "provinceCode": "28" },\n'
    '  …\n'
    ']'
)
doc.add_paragraph('Shape con slug: objeto de data/cuadro-medico/provincias/<slug>.json — '
                  '{ slug, displayName, provinceCode, especialidades: [<spec-slugs>] }.')

# --- api/especialidades.js ---
doc.add_heading('Endpoint 6 · api/especialidades.js (listado / detalle de especialidad)', level=4)
doc.add_paragraph('Ruta: api/especialidades.js · 36 líneas · HTTP: GET /api/especialidades')
doc.add_paragraph('Query: slug (opcional).')
doc.add_paragraph('Cachés: masterCache.')
doc.add_paragraph('Errores: 404 "Especialidad no encontrada", 404 "Sin datos para".')
doc.add_paragraph('Shape sin slug:')
add_code_block(
    '[\n'
    '  { "slug": "cardiologia", "name": "Cardiología",\n'
    '    "nameApi": "CARDIOLOGÍA",\n'
    '    "professionalPlural": "Cardiólogos",\n'
    '    "professionalPluralLower": "cardiólogos",\n'
    '    "kind": "specialty" | "service" | "technique",\n'
    '    "specialityCode": 9 },\n'
    '  …\n'
    ']'
)
doc.add_paragraph('Shape con slug: meta + array provincias[]:')
add_code_block(
    '{\n'
    '  "slug": "cardiologia", "name": "Cardiología", "kind": "specialty",\n'
    '  "provincias": [\n'
    '    { "slug": "madrid", "displayName": "Madrid", "count": 381 },\n'
    '    …  // ordenado por count desc\n'
    '  ]\n'
    '}'
)

# --- api/providers-detail.js ---
doc.add_heading('Endpoint 7 · api/providers-detail.js (detalle por locCode)', level=4)
doc.add_paragraph('Ruta: api/providers-detail.js · 78 líneas · HTTP: GET /api/providers-detail')
doc.add_paragraph('Query: id (String, obligatorio) — providerLocalicationCode.')
doc.add_paragraph('Lee data/provider-details/<id>.json. Sin caché in-memory.')
doc.add_paragraph(
    'Un mismo provider-details puede tener varios entries (mismo provider, distintas '
    'especialidades). mapDetail() consolida: toma base del primer entry, deduplica y agrupa '
    'specialities y languages de todos los entries.'
)
doc.add_paragraph('Consumido por: usado como fallback / enriquecimiento desde otros endpoints; '
                  'también accesible directo desde el frontend si se necesita una sola ubicación.')

# --- api/specialities.js ---
doc.add_heading('Endpoint 8 · api/specialities.js (proxy autocomplete a ASISA)', level=4)
doc.add_paragraph('Ruta: api/specialities.js · 32 líneas · HTTP: GET /api/specialities')
doc.add_paragraph('Query: provinceCode (opcional, código numérico INE).')
doc.add_paragraph('NO lee data/ ni tiene caché — es un PROXY EN VIVO al backend de ASISA:')
add_code_block(
    'GET https://ursaepre.asisa.es/ASISA/middlewasisa/public/v1/api/searchPortal/\n'
    '    autocomplete/specialities\n'
    '  ?specialityDescription=&networkCode=1[&provinceCode=<code>]\n'
    '  &maxResultsNumber=500\n'
    'Headers:\n'
    '  Ocp-Apim-Subscription-Key: 0908b85b9d0e4a75b2eb33048bd9fe01\n'
    '  Api-Version: 1'
)
doc.add_paragraph('Devuelve la respuesta tal cual (passthrough). Si el backend de ASISA cae, '
                  'este endpoint cae también. Pensado para autocomplete de búsqueda.')
doc.add_paragraph('CRÍTICO para migración Azure: este endpoint es el único que sí depende del '
                  'backend ASISA en runtime. La subscription key debe ir a Key Vault.')

# --- api/sitemap.js + api/sitemap-cuadro-medico.js ---
doc.add_heading('Endpoints 9 y 10 · sitemaps (Vercel-only fallback)', level=4)
doc.add_paragraph(
    'api/sitemap.js (29 líneas) → /sitemap.xml: sitemap index estático que apunta a 5 '
    'sitemaps específicos en www.asisa.es. Exporta getSitemapIndexXml() para que api/markup.js '
    'lo reuse.'
)
doc.add_paragraph(
    'api/sitemap-cuadro-medico.js (94 líneas) → /sitemap-cuadro-medico-<type>.xml: genera el '
    'XML del tipo según el query param type. Tipos aceptados: provincias, provincia-specs, '
    'doctores, centros, especialidades. Cada uno enumera URLs leyendo:'
)
add_bullet('provincias → data/cuadro-medico/provincias/*.json (52 URLs)')
add_bullet('provincia-specs → cruza cada provincia con sus especialidades (~3.250 URLs)')
add_bullet('doctores → keys de doctores-index.json (~20.500 URLs)')
add_bullet('centros → keys de centros-index.json (~6.500 URLs)')
add_bullet('especialidades → ficheros en especialidades/ (~181 URLs)')
doc.add_paragraph(
    'Estos endpoints existen como vía paralela. EDS sirve los mismos paths NATIVAMENTE vía '
    'helix-sitemap.yaml (con origin: https://www.asisa.es); los de Vercel son un fallback '
    'accesible directamente en asisa-pc.vercel.app y útil para debugging.'
)

# --- api/sync-aem.js ---
doc.add_heading('Endpoint 11 · api/sync-aem.js (sincronizador de sitemap externo)', level=4)
doc.add_paragraph('Ruta: api/sync-aem.js · 84 líneas · HTTP: GET /api/sync-aem (interno)')
doc.add_paragraph('Query:')
add_bullet('secret (String) — debe coincidir con env SYNC_SECRET, si no 401.')
add_bullet('path (String, opcional) — sincroniza solo esa ruta.')
add_bullet('limit (Number, default 50) — URLs por lote.')
add_bullet('offset (Number, default 0) — offset para paginación.')
doc.add_paragraph('Si no se pasa path, lee https://www.asisa.es/sitemap.xml, filtra URLs que '
                  'contengan "/cuadro-medico/" y procesa por lotes. Para cada URL hace:')
add_code_block(
    'POST https://admin.hlx.page/preview/asisa-softtek/asisa-pc/main<path>\n'
    'POST https://admin.hlx.page/live/asisa-softtek/asisa-pc/main<path>\n'
    'Headers: x-auth-token: <HLX_ADMIN_API_TOKEN>'
)
doc.add_paragraph('Concurrencia 5 (Promise.all en lotes). Requiere env HLX_ADMIN_API_TOKEN o '
                  'devuelve 500. Útil para automatizar sincronización masiva desde un sitemap '
                  'externo (e.g. tras una republicación en www.asisa.es).')

doc.add_heading('3.4.2 Resumen de cachés y TTLs', level=3)
add_table(
    ['Endpoint', 'Cachés in-memory', 'Cache-Control'],
    [
        ('api/markup.js', '—', 'max-age=60'),
        ('api/providers.js',
         'provinciasCache + especialidadesCache + allProvinceCache',
         's-maxage=300, stale-while-revalidate=3600'),
        ('api/doctor.js', 'indexCache + providersListCache',
         's-maxage=3600, stale-while-revalidate=86400'),
        ('api/centro.js',
         'centrosIndexCache + doctoresIndexCache + provinceScanCache',
         's-maxage=3600, stale-while-revalidate=86400'),
        ('api/provincias.js', '—',
         's-maxage=86400, stale-while-revalidate=604800'),
        ('api/especialidades.js', 'masterCache',
         's-maxage=86400, stale-while-revalidate=604800'),
        ('api/providers-detail.js', '—',
         's-maxage=3600, stale-while-revalidate=86400'),
        ('api/specialities.js', '— (proxy en vivo)',
         's-maxage=86400, stale-while-revalidate=604800'),
        ('api/sitemap.js', '—', 's-maxage=3600, stale-while-revalidate=86400'),
        ('api/sitemap-cuadro-medico.js', '—',
         's-maxage=86400, stale-while-revalidate=86400'),
        ('api/sync-aem.js', '— (no cacheable)', '—'),
    ],
    widths=[Inches(2.2), Inches(2.5), Inches(2)],
)
doc.add_paragraph(
    'Nota Azure: las cachés in-memory SOBREVIVEN dentro de la misma instancia warm pero NO entre '
    'instancias ni tras cold start. En Premium Plan con Always-On las cachés serán efectivas la '
    'mayoría del tiempo. Si quieres caché compartida, Azure Cache for Redis.'
)

doc.add_heading('3.5 Bloques EDS (carpeta blocks/)', level=2)
doc.add_paragraph(
    'Cada bloque vive en blocks/<nombre>/<nombre>.{js,css}. EDS los instancia automáticamente '
    'cuando detecta <div class="<nombre>"> en el HTML. Los CSS son casi todos vacíos: los estilos '
    'efectivos vienen del design system de ASISA (clientlib-site.min.css).'
)
doc.add_paragraph('Bloques heredados del boilerplate AEM:')
add_bullet('cards, columns, hero, fragment — contenido estático autorado en AEM.')
add_bullet('header — carga /nav como fragmento. Hamburguesa responsive, dropdowns, breakpoint 900px.')
add_bullet('footer — carga /footer como fragmento (.plain.html).')

doc.add_paragraph('Bloques específicos del cuadro médico (todos operan en modo BYOM: leen '
                  'window.location.pathname, extraen slugs, llaman a las APIs):')
add_table(
    ['Bloque', 'URL donde aparece', 'APIs que consume'],
    [
        ('cuadro-medico', '/p/<prov>, /p/<prov>/pe/<spec>, /e/<spec>',
         'providers, provincias, especialidades'),
        ('cuadro-medico-provincias', '/cuadro-medico (home)',
         'provincias'),
        ('cuadro-medico-top-especialidades', '/cuadro-medico (home)',
         'especialidades'),
        ('cuadro-medico-otras-especialidades', '/p/<prov>/pe/<spec>, /e/<spec>',
         'provincias, especialidades'),
        ('cuadro-medico-otras-provincias', '/p/<prov>/pe/<spec>, /e/<spec>',
         'especialidades'),
        ('cuadro-medico-ficha-doctor', '/d/<key>',
         'doctor, provincias'),
        ('cuadro-medico-otros-medicos', '/d/<key>',
         'doctor, provincias, providers'),
        ('cuadro-medico-ficha-centro', '/c/<key>',
         'centro, provincias'),
    ],
    widths=[Inches(2.2), Inches(2.5), Inches(1.8)],
)
doc.add_paragraph(
    'Regla de oro de los bloques: TODAS las llamadas a /api/* deben usar URL ABSOLUTA del overlay '
    '("https://<overlay-host>/api/...") porque el bloque corre en aem.live, no en Vercel/Azure. '
    'Una llamada relativa fetch("/api/…") apuntaría al dominio de EDS, que no tiene esos '
    'endpoints. En la migración a Azure este punto es crítico: hay que cambiar la base URL en '
    'todos los bloques.'
)

doc.add_page_break()
doc.add_heading('3.5.1 Detalle exhaustivo de cada bloque del cuadro médico', level=3)
doc.add_paragraph(
    'A continuación se documenta cada bloque del cuadro médico fichero por fichero: ruta, '
    'propósito, URLs donde aparece, cómo lee el contexto (URL o filas del DOM), endpoints API '
    'consumidos con sus query params exactos, esqueleto del HTML que pinta, estados especiales '
    '(ocultado, error, loading), y funciones auxiliares relevantes. Los bloques boilerplate '
    '(cards, columns, header, footer, fragment, hero) tienen un resumen al final.'
)
doc.add_paragraph(
    'Constantes globales que comparten todos los bloques:'
)
add_table(
    ['Constante', 'Valor', 'Para qué'],
    [
        ('API_BASE', 'https://asisa-pc.vercel.app', 'Base URL de las APIs de datos'),
        ('ASISA_SEARCH_PRIVATE', 'https://www.asisa.es/asegurado/salud/cuadro-medico/'
                                  'resultados-cuadro-medico',
         'Destino "Pedir cita" para asegurados (área privada)'),
        ('ASISA_SEARCH_PUBLIC', 'https://www.asisa.es/cuadro-medico/resultados-cuadro-medico',
         'Destino "Compartir" público'),
        ('PAGE_SIZE', '10', 'Resultados por página en el bloque cuadro-medico'),
        ('networkId / specialityType', "'1'", 'Hardcoded en params de Pedir cita (red Salud)'),
        ('Breakpoint responsive', '900px',
         'Usado por header para hamburger menu (window.matchMedia)'),
    ],
    widths=[Inches(2), Inches(2.5), Inches(2)],
)

# -- BLOQUE 1: cuadro-medico --
doc.add_heading('Bloque 1 · cuadro-medico (motor de búsqueda principal)', level=4)
doc.add_paragraph('Ruta: blocks/cuadro-medico/cuadro-medico.js (274 líneas).')
doc.add_paragraph(
    'Propósito: motor principal de listado paginado de profesionales y centros, con tabs y '
    'paginación. Es el bloque más grande y complejo del sitio.'
)
doc.add_paragraph('URLs donde aparece:')
add_bullet('/cuadro-medico/p/<prov>')
add_bullet('/cuadro-medico/p/<prov>/pe/<spec>')
add_bullet('/cuadro-medico/e/<spec> (búsqueda nacional por especialidad)')
doc.add_paragraph(
    'Lee el contexto desde window.location.pathname con getSlugsFromUrl(): extrae provSlug '
    '(segmento "p/"), specSlug (segmento "pe/" o "e/") y un flag nationalSpec (true si hay '
    'especialidad pero no provincia). Si no hay provSlug ni nationalSpec, el bloque se oculta.'
)
doc.add_paragraph('APIs consumidas:')
add_table(
    ['Endpoint', 'Query params', 'Para qué'],
    [
        ('/api/provincias', 'slug=<provSlug>', 'displayName y provinceCode'),
        ('/api/especialidades', 'slug=<specSlug>', 'name de la especialidad'),
        ('/api/providers',
         'tab=professionals|centers, page=N, limit=10, provinceSlug, specSlug',
         'listado paginado + totales por tab'),
    ],
    widths=[Inches(1.6), Inches(2.4), Inches(2.5)],
)
doc.add_paragraph('Estructura DOM que pinta (esqueleto):')
add_code_block(
    '<div class="cmp-medical-picture-result">\n'
    '  <section class="eds-mp-box-head">\n'
    '    <h1 class="eds-mp-box-head--title">…</h1>\n'
    '    <p class="eds-mp-box-head--text">…</p>\n'
    '  </section>\n'
    '  <div class="eds-mp-tabs">\n'
    '    <ul class="eds-mp-tabs__nav">\n'
    '      <li class="eds-mp-tabs__nav--item active" data-tab="professionals">\n'
    '         Profesionales (X)</li>\n'
    '      <li class="eds-mp-tabs__nav--item" data-tab="centers">Centros (Y)</li>\n'
    '    </ul>\n'
    '    <div class="eds-mp-tabs__container">\n'
    '      <div class="eds-mp-tabs__content">\n'
    '        <div class="eds-mp-card">…</div>  <!-- una por resultado -->\n'
    '      </div>\n'
    '      <div class="eds-mp-pagination">…</div>\n'
    '    </div>\n'
    '  </div>\n'
    '</div>'
)
doc.add_paragraph('Cada tarjeta de resultado tiene esta estructura:')
add_code_block(
    '<div class="eds-mp-card">\n'
    '  <div class="eds-mp-card__principal-tag">\n'
    '    <div class="cmp-tag-template cmp-tag-template--blue">MÉDICO|HOSPITAL|CENTRO|…</div>\n'
    '    <!-- tags condicionales: Centro de ASISA, Receta electrónica -->\n'
    '  </div>\n'
    '  <div class="eds-mp-card__info">\n'
    '    <div class="eds-mp-card__info--contact">\n'
    '      <p class="eds-mp-card__type--speciality">…</p>\n'
    '      <p class="eds-mp-card__type--name">Dr. Nombre</p>\n'
    '      <p class="eds-mp-card__type--num-member">Núm. Colegiado</p>\n'
    '      <p class="eds-mp-card__type--center">Centro padre</p>\n'
    '      <div class="eds-mp-card__type--address">Dirección</div>\n'
    '      <div class="eds-mp-card__info--location">\n'
    '        <a href="https://maps.google.com/…">Cómo llegar</a>\n'
    '        <a href="tel:…">Teléfono</a>\n'
    '      </div>\n'
    '    </div>\n'
    '    <div class="eds-mp-card__info--tags">\n'
    '      <!-- Cita online, idiomas -->\n'
    '    </div>\n'
    '  </div>\n'
    '  <div class="eds-mp-card__info--buttons">\n'
    '    <a href="<ASISA_SEARCH_PRIVATE>?…">Pedir cita</a>\n'
    '    <a href="/cuadro-medico/d/<detailUrl>">Ver detalle</a>\n'
    '  </div>\n'
    '</div>'
)
doc.add_paragraph('Funciones auxiliares:')
add_bullet('getSlugsFromUrl(): parsea pathname y devuelve {provSlug, specSlug, nationalSpec}.')
add_bullet('buildCitaUrl(): construye URL a ASISA_SEARCH_PRIVATE con networkId=1, '
           'specialityType=1, lat, lon y código de provincia.')
add_bullet('formatPersonName(): mapea "APELLIDOS, NOMBRE" → "Dr./Dra. Nombre Apellidos" según '
           'terminación del primer nombre.')
add_bullet('getProviderTag(): mapea doctorType y providerType al texto de la etiqueta principal '
           '(3=HOSPITAL, 4=CENTRO MÉDICO, 8=LABORATORIO, 2=TRANSPORTE, 9=OXIGENOTERAPIA).')
add_bullet('renderCard(), renderPagination(): generación de HTML.')
add_bullet('attachListeners(): clicks en tabs y páginas refetchean conservando el otro estado.')
doc.add_paragraph('Estados especiales:')
add_bullet('Sin provSlug ni nationalSpec → block.hidden = true.')
add_bullet('API devuelve 0 resultados en TODAS las categorías → block.hidden = true.')
add_bullet('Error de fetch → muestra <div class="cmp-medical-picture-result__error">.')
add_bullet('Loading → spinner inline hasta la primera respuesta.')

add_para('Reglas de negocio del listado paginado (caps):', bold=True)
add_table(
    ['Parámetro', 'Valor', 'Definido en', 'Por qué'],
    [
        ('DEFAULT_LIMIT', '10',
         'api/providers.js:4',
         'Tamaño de página por defecto cuando el bloque no pasa limit explícito.'),
        ('MAX_LIMIT', '50',
         'api/providers.js:5',
         'Tope al limit que un cliente puede pedir en una sola request (evita XHR enormes).'),
        ('MAX_TOTAL_BY_TAB.professionals', '30',
         'api/providers.js:6',
         'TOTAL de profesionales mostrables en el cuadro (≈ 3 páginas de 10). El backend trunca '
         'la lista a 30 antes de paginar — el resto NO es alcanzable navegando.'),
        ('MAX_TOTAL_BY_TAB.centers', '50',
         'api/providers.js:6',
         'TOTAL de centros mostrables (≈ 5 páginas de 10).'),
    ],
    widths=[Inches(1.8), Inches(0.7), Inches(1.4), Inches(2.6)],
)
doc.add_paragraph(
    'Justificación: el backend de ASISA puede devolver listados muy largos (cientos de médicos '
    'por provincia+especialidad). Mostrarlos todos haría la página injustificadamente larga y '
    'mal indexable. La regla de producto acuerda topes razonables (30 profesionales, 50 centros) '
    'y obliga al usuario a refinar por provincia/especialidad/centro para encontrar lo que busca. '
    'La etiqueta del tab refleja el conteo CAPADO (e.g. "Profesionales (30)") para no engañar '
    'sobre lo navegable.'
)
doc.add_paragraph(
    'Implicación en pagination: con tab=professionals y limit=10, totalPages será como máximo 3 '
    '(30/10). Con tab=centers, como máximo 5 (50/10). El bloque renderiza paginación con elipsis '
    'si hace falta, pero en la práctica nunca pasa de 5 páginas.'
)
doc.add_paragraph(
    'Si en el futuro hay que cambiar estos caps (p.ej. para SEO o por petición de marketing), '
    'es un cambio de UNA línea en api/providers.js. El frontend no asume nada del cap: respeta '
    'totalPages que viene en la respuesta.'
)
doc.add_paragraph(
    'Otros usos del endpoint: el bloque cuadro-medico-otros-medicos pide '
    'limit=50&tab=professionals para extraer hasta 50 chips. Como MAX_LIMIT=50, no hay que tocar '
    'caps. Si se necesitara más, primero hay que subir MAX_LIMIT.'
)

# -- BLOQUE 2: cuadro-medico-provincias --
doc.add_heading('Bloque 2 · cuadro-medico-provincias (selector de provincias)', level=4)
doc.add_paragraph('Ruta: blocks/cuadro-medico-provincias/cuadro-medico-provincias.js (28 líneas).')
doc.add_paragraph('Propósito: lista todas las provincias donde ASISA tiene red.')
doc.add_paragraph('URLs donde aparece: home y landings de cuadro médico (autoradas en AEM).')
doc.add_paragraph('Sin contexto de entrada — modo BYOM puro (no lee URL).')
doc.add_paragraph('API: GET /api/provincias (sin params) → array {slug, displayName}.')
add_code_block(
    '<h2 class="eds-md-prov-title">Provincias donde está ASISA (X)</h2>\n'
    '<ul class="eds-md-prov-list">\n'
    '  <li class="eds-md-prov-item">\n'
    '    <a href="/cuadro-medico/p/{slug}">\n'
    '      <span><i class="icon-localizacion"></i></span>\n'
    '      <p>{displayName}</p>\n'
    '    </a>\n'
    '  </li>\n'
    '</ul>'
)
doc.add_paragraph('Estados: spinner durante el fetch; error muestra '
                  '"No se pudieron cargar las provincias".')

# -- BLOQUE 3: cuadro-medico-top-especialidades --
doc.add_heading('Bloque 3 · cuadro-medico-top-especialidades', level=4)
doc.add_paragraph('Ruta: blocks/cuadro-medico-top-especialidades/cuadro-medico-top-especialidades.js '
                  '(28 líneas).')
doc.add_paragraph('Propósito: lista de especialidades más buscadas con link a la búsqueda nacional.')
doc.add_paragraph('URLs donde aparece: home / landings de cuadro médico.')
doc.add_paragraph('API: GET /api/especialidades → array {slug, name}.')
add_code_block(
    '<h2 class="eds-md-esp-top-title">Especialidades más buscadas (X)</h2>\n'
    '<ul class="eds-md-esp-top-list">\n'
    '  <li class="eds-md-esp-top-item">\n'
    '    <a href="/cuadro-medico/e/{slug}">\n'
    '      <span><i class="icon-ventajas"></i></span>\n'
    '      <p>{name}</p>\n'
    '    </a>\n'
    '  </li>\n'
    '</ul>'
)

# -- BLOQUE 4: cuadro-medico-otras-especialidades --
doc.add_heading('Bloque 4 · cuadro-medico-otras-especialidades', level=4)
doc.add_paragraph('Ruta: blocks/cuadro-medico-otras-especialidades/cuadro-medico-otras-especialidades.js '
                  '(80 líneas).')
doc.add_paragraph(
    'Propósito: chips con especialidades alternativas. Dos modos según la URL:'
)
add_bullet('Provincia + especialidad (/cuadro-medico/p/<prov>/pe/<spec>): muestra las otras '
           'especialidades disponibles EN ESA PROVINCIA.')
add_bullet('Especialidad nacional (/cuadro-medico/e/<spec>): muestra TOP 15 especialidades '
           'nacionales (filtra las de kind === "service").')
doc.add_paragraph(
    'Lee de la URL: getSlugsFromUrl() devuelve provSlug (segmento "p/") y specSlug (segmento '
    '"pe/" o "e/"). Sin ninguno de los dos → block.hidden = true.'
)
doc.add_paragraph('APIs:')
add_table(
    ['Modo', 'Endpoints'],
    [
        ('Provincial',
         'GET /api/provincias?slug=<provSlug> + GET /api/especialidades → filtra por '
         'provincia.especialidades[]'),
        ('Nacional', 'GET /api/especialidades → toma top 15 con kind !== "service"'),
    ],
    widths=[Inches(1.5), Inches(5)],
)
doc.add_paragraph('Estructura DOM (provincial):')
add_code_block(
    '<div class="eds-mp-other-specs">\n'
    '  <h2 class="eds-mp-other-specs__title">Otras especialidades en {provincia}</h2>\n'
    '  <ul class="eds-mp-other-specs__container">\n'
    '    <li><a class="cmp-tag-template cmp-tag-template--blue-100"\n'
    '           href="/cuadro-medico/p/{provSlug}/pe/{slug}">\n'
    '      <span class="cmp-tag-template__text">{name}</span></a></li>\n'
    '  </ul>\n'
    '</div>'
)
doc.add_paragraph(
    'En modo nacional usa clases cmp-medical-detail__subtitle y cmp-medical-detail__'
    'other-specialities, y marca la especialidad actual con cmp-tag-template--blue.'
)

# -- BLOQUE 5: cuadro-medico-otras-provincias --
doc.add_heading('Bloque 5 · cuadro-medico-otras-provincias', level=4)
doc.add_paragraph('Ruta: blocks/cuadro-medico-otras-provincias/cuadro-medico-otras-provincias.js '
                  '(47 líneas).')
doc.add_paragraph('Propósito: grid de cards de OTRAS provincias que ofrecen la misma especialidad. '
                  'Cada card muestra el conteo de profesionales y enlaza a la combinación '
                  '/p/<slug>/pe/<spec>.')
doc.add_paragraph(
    'URLs donde aparece: /cuadro-medico/p/<prov>/pe/<spec> y /cuadro-medico/e/<spec>. El bloque '
    'lee specSlug tanto del segmento "pe/" como del "e/". Sin specSlug → block.hidden = true.'
)
doc.add_paragraph('API: GET /api/especialidades?slug=<specSlug> → especialidad con array provincias[].')
add_code_block(
    '<h2 class="cm-otras-prov-title">Otras provincias con {specName} ASISA</h2>\n'
    '<div class="cm-otras-prov-list">\n'
    '  <article class="cm-otras-prov-card">\n'
    '    <h3 class="cm-otras-prov-card__name">{specName} {provincia}</h3>\n'
    '    <p class="cm-otras-prov-card__count">{count} profesionales</p>\n'
    '    <a class="cm-otras-prov-card__arrow"\n'
    '       href="/cuadro-medico/p/{slug}/pe/{specSlug}">→</a>\n'
    '  </article>\n'
    '</div>'
)
doc.add_paragraph(
    'Excluye la provincia actual del listado (filter p.slug !== provSlug). En la búsqueda '
    'nacional /e/<spec> muestra TODAS las provincias con la especialidad.'
)

# -- BLOQUE 6: cuadro-medico-ficha-doctor --
doc.add_heading('Bloque 6 · cuadro-medico-ficha-doctor (ficha del profesional)', level=4)
doc.add_paragraph('Ruta: blocks/cuadro-medico-ficha-doctor/cuadro-medico-ficha-doctor.js (199 líneas).')
doc.add_paragraph(
    'Propósito: ficha completa de un médico (agrupado por número colegiado). Como un médico puede '
    'trabajar en varios centros, renderiza header + una "card de ubicación" + CTA por cada centro '
    'donde ejerce.'
)
doc.add_paragraph('URLs: /cuadro-medico/d/<key>. getKeyFromUrl() extrae el segmento "d/".')
doc.add_paragraph('APIs:')
add_table(
    ['Endpoint', 'Para qué'],
    [
        ('/api/doctor?key=<key>',
         'doctor completo con name, collegiateCode, specialities[], locations[]'),
        ('/api/provincias', 'mapa provinceSlug → displayName'),
    ],
    widths=[Inches(2.2), Inches(4.3)],
)
doc.add_paragraph('Esqueleto DOM (resumido):')
add_code_block(
    '<div class="cmp-medical-detail">\n'
    '  <section class="eds-mp-box-head">  <!-- H1 + intro SEO -->\n'
    '    <h1>Dr. Nombre, Especialidad</h1>\n'
    '  </section>\n'
    '  <div class="cmp-medical-detail__first-block">  <!-- primer centro -->\n'
    '    <div class="cmp-medical-detail__title-block">…</div>\n'
    '    <div class="cmp-medical-detail__address-block">…</div>\n'
    '    <div class="cmp-medical-detail__buttons-block">…</div>\n'
    '  </div>\n'
    '  <div class="cmp-medical-detail__first-block cmp-medical-detail--blue">\n'
    '    <h2>{especialidad}</h2>\n'
    '    <a href="<cita>">Pedir cita online</a>\n'
    '  </div>\n'
    '  <!-- repite el bloque por cada ubicación adicional -->\n'
    '  <h2 class="cmp-medical-detail__subtitle">\n'
    '     Dr. Nombre también trabaja en estos centros\n'
    '  </h2>\n'
    '</div>'
)
doc.add_paragraph('Funciones clave: buildShareUrl() (URL pública sin fromPublicArea) y '
                  'buildCitaUrl() (URL privada con fromPublicArea=true).')
doc.add_paragraph('Estados: sin key → block.hidden; sin locations o error → '
                  '"No se pudo cargar la ficha del médico".')

# -- BLOQUE 7: cuadro-medico-otros-medicos --
doc.add_heading('Bloque 7 · cuadro-medico-otros-medicos', level=4)
doc.add_paragraph('Ruta: blocks/cuadro-medico-otros-medicos/cuadro-medico-otros-medicos.js (86 líneas).')
doc.add_paragraph(
    'Propósito: bajo la ficha del doctor, dos listas de chips con OTROS médicos: (1) de la misma '
    'especialidad en la misma provincia; (2) de la misma especialidad en el mismo centro padre '
    '(solo si el doctor tiene parentDescription).'
)
doc.add_paragraph('URLs: /cuadro-medico/d/<key>.')
doc.add_paragraph('APIs:')
add_table(
    ['Endpoint', 'Para qué'],
    [
        ('/api/doctor?key=<key>',
         'datos del doctor (specSlug, provinceSlug, parentDescription, name)'),
        ('/api/provincias?slug=<provinceSlug>', 'displayName'),
        ('/api/providers?provinceSlug&specSlug&tab=professionals&limit=50',
         'listado para extraer chips (máx 20 en provincia, máx 10 en mismo centro)'),
    ],
    widths=[Inches(3.2), Inches(3.3)],
)
add_code_block(
    '<section>\n'
    '  <h2 class="cmp-medical-detail__subtitle">\n'
    '    Otros médicos de {especialidad} en {provincia}\n'
    '  </h2>\n'
    '  <div class="cmp-medical-detail__other-specialities">\n'
    '    <a class="cmp-tag-template cmp-tag-template--blank"\n'
    '       href="/cuadro-medico/d/{detailUrl}">\n'
    '      <span class="cmp-tag-template__text">Dr. Nombre</span></a>\n'
    '  </div>\n'
    '</section>\n'
    '<!-- segunda sección equivalente para "en {centro}" si aplica -->'
)
doc.add_paragraph('Si las dos listas vienen vacías → block.hidden = true.')

# -- BLOQUE 8: cuadro-medico-ficha-centro --
doc.add_heading('Bloque 8 · cuadro-medico-ficha-centro (ficha del centro)', level=4)
doc.add_paragraph('Ruta: blocks/cuadro-medico-ficha-centro/cuadro-medico-ficha-centro.js (277 líneas).')
doc.add_paragraph(
    'Propósito: ficha completa de un centro médico. Incluye breadcrumb, card principal del centro, '
    'sección con acordeones por especialidad (con doctores, subespecialidades y observaciones), '
    'grid de doctores del centro y sección de "otros centros ASISA con las mismas especialidades '
    'en la provincia".'
)
doc.add_paragraph('URLs: /cuadro-medico/c/<key>.')
doc.add_paragraph('APIs:')
add_table(
    ['Endpoint', 'Para qué'],
    [
        ('/api/centro?key=<key>',
         'centro completo con specialities[], doctors[], otherCentros[]'),
        ('/api/provincias', 'mapeo provinceSlug → displayName'),
    ],
    widths=[Inches(2.2), Inches(4.3)],
)
doc.add_paragraph('Esqueleto DOM (resumido):')
add_code_block(
    '<div class="cmp-medical-detail">\n'
    '  <nav class="cmp-breadcrumb">…</nav>\n'
    '  <section class="eds-mp-box-head"><h1>…</h1></section>\n'
    '  <div class="cmp-medical-detail__first-block">\n'
    '    <!-- card principal del centro: tags, título, dirección, teléfono -->\n'
    '  </div>\n'
    '  <section class="cm-fcentro__specs-section">\n'
    '    <h2 class="cmp-medical-detail__subtitle">Especialidades del centro</h2>\n'
    '    <div class="cm-fcentro__spec">\n'
    '      <div class="cm-fcentro__spec-header">\n'
    '        <h3>Dermatología</h3>\n'
    '        <a href="tel:">Teléfono</a>  <a href="<cita>">Pedir cita</a>\n'
    '      </div>\n'
    '      <details class="cm-fcentro__spec-details">\n'
    '        <summary>Ver más información</summary>\n'
    '        <div class="cm-fcentro__spec-body">\n'
    '          <!-- columnas: doctores, subespecialidades, observaciones -->\n'
    '        </div>\n'
    '      </details>\n'
    '    </div>\n'
    '  </section>\n'
    '  <section class="cm-fcentro__doctors-section">\n'
    '    <div class="cm-fcentro__doctors-grid">\n'
    '      <article class="cm-fcentro__doctor-card">\n'
    '        <i class="icon-personal-asisa-{mujer|hombre}"></i>\n'
    '        <h3>Dr. Nombre</h3><p>Especialidad</p>\n'
    '        <a href="/cuadro-medico/d/{key}">Ver perfil</a>\n'
    '      </article>\n'
    '    </div>\n'
    '  </section>\n'
    '  <section class="cm-fcentro__other-section">\n'
    '    <h2>Otros centros ASISA con las mismas especialidades</h2>\n'
    '    <!-- grid de otros centros -->\n'
    '  </section>\n'
    '</div>'
)
doc.add_paragraph('Funciones notables:')
add_bullet('doctorIconClass(d): elige el icono (mujer/hombre) según gender o heurística por "Dra.".')
add_bullet('showsPedirCita(spec): false si la especialidad contiene "urgenc" (no se puede pedir '
           'cita para urgencias). True si hay doctors o cita online.')
add_bullet('Cada sección se omite si no hay datos: si no hay doctors, no se renderiza la grid; '
           'si no hay otros centros, no aparece la sección.')

# -- Boilerplate brief --
doc.add_heading('Bloques boilerplate (resumen)', level=4)
add_bullet('cards: transforma rows del editor en grid <ul><li>. Optimiza imágenes con '
           'createOptimizedPicture (ancho 750px).')
add_bullet('columns: aplica clase .columns-N-cols dinámica y .columns-img-col para columnas que '
           'solo contienen una imagen.')
add_bullet('header: carga /nav como fragmento. Hamburger menu, ARIA, dropdowns con aria-expanded, '
           'breakpoint 900px.')
add_bullet('footer: carga /footer como fragmento.')
add_bullet('fragment: utility para incluir HTML de otra ruta .plain.html. Resetea media URLs.')
add_bullet('hero: stub vacío (no implementado).')

doc.add_heading('Reglas de negocio consolidadas (caps y umbrales)', level=4)
doc.add_paragraph(
    'Esta tabla consolida todos los "magic numbers" del proyecto. Son acuerdos de producto o '
    'topes técnicos que el equipo nuevo debe conocer: tocar uno puede cambiar visiblemente la '
    'experiencia del usuario o la cobertura SEO.'
)
add_table(
    ['Cap / umbral', 'Valor', 'Definido en', 'Qué controla'],
    [
        ('Tamaño de página por defecto', '10',
         'api/providers.js:4 (DEFAULT_LIMIT)',
         'Resultados por página cuando el bloque no pasa limit.'),
        ('Tamaño máximo de página', '50',
         'api/providers.js:5 (MAX_LIMIT)',
         'Tope al limit por request. Si un cliente pide más, se trunca a 50.'),
        ('Total profesionales mostrables', '30',
         'api/providers.js:6 (MAX_TOTAL_BY_TAB.professionals)',
         'Tope total del listado en el tab "Profesionales" (≈ 3 páginas de 10). El backend trunca '
         'la lista antes de paginar; el resto NO es alcanzable.'),
        ('Total centros mostrables', '50',
         'api/providers.js:6 (MAX_TOTAL_BY_TAB.centers)',
         'Tope total en el tab "Centros médicos" (≈ 5 páginas de 10).'),
        ('Top especialidades nacionales', '15',
         'blocks/cuadro-medico-otras-especialidades/cuadro-medico-otras-especialidades.js:41',
         'Especialidades mostradas como chips en /cuadro-medico/e/<spec> (filtra '
         'kind !== "service" antes).'),
        ('Otros médicos en provincia (chips)', '20',
         'blocks/cuadro-medico-otros-medicos/cuadro-medico-otros-medicos.js:78',
         'Chips máximos en la lista "Otros médicos de {spec} en {provincia}" bajo la ficha de '
         'doctor.'),
        ('Otros médicos en centro (chips)', '10',
         'blocks/cuadro-medico-otros-medicos/cuadro-medico-otros-medicos.js:79',
         'Chips máximos en "Otros médicos de {spec} en {centro}" — solo aparece si el doctor '
         'tiene parentDescription.'),
        ('Otros centros con mismas especialidades', '— (configurable)',
         'api/centro.js:141',
         'Centros similares mostrados al pie de la ficha de centro. El cap se pasa como argumento '
         'a la función; revisa el código si se va a modificar.'),
        ('Especialidades visibles por centro', '4 (resto colapsadas)',
         'api/centro.js:145',
         'Solo las primeras 4 especialidades se muestran expandidas por defecto en el card de un '
         'centro; el resto aparecen colapsadas tras un toggle.'),
        ('Concurrencia de fetch a API ASISA',
         '10 (providers), 25 (provider-details)',
         'generate-providers-data.mjs y generate-provider-details.mjs',
         'Requests paralelas a /searchPortal y /providers/details. Bajar a 5/10 si aparecen 429.'),
        ('Concurrencia de admin EDS', '10',
         'refresh-eds-pages.mjs (CONCURRENCY)',
         'Requests paralelas a admin.hlx.page (preview/live/code/index). Subir podría disparar '
         'rate limits.'),
        ('Concurrencia de copy AEM Author', '5',
         'create-aem-pages.mjs (CONCURRENCY)',
         'Copias paralelas vía /bin/wcmcommand. Más causa OakLock collisions; menos enlentece.'),
        ('Cache TTL CDN', 's-maxage=300, stale-while-revalidate=3600',
         'api/providers.js, api/doctor.js, api/centro.js (res.setHeader)',
         '5 min de cache fresh + 1h de stale en el edge. Tras un update de datos, los listados '
         'pueden tardar hasta 5 min en reflejar.'),
        ('Cache TTL sitemaps',
         's-maxage=86400, stale-while-revalidate=86400',
         'api/sitemap.js, api/sitemap-cuadro-medico.js',
         '24h de cache fresh + 24h stale para sitemaps.'),
    ],
    widths=[Inches(1.6), Inches(0.9), Inches(2), Inches(2)],
)

doc.add_heading('Regla de oro común a todos los bloques', level=4)
doc.add_paragraph(
    'TODAS las llamadas a /api/* DEBEN usar URL ABSOLUTA "https://<overlay-host>/api/…". Una '
    'llamada relativa fetch("/api/…") apunta al dominio de aem.live, que no tiene los endpoints '
    'y devuelve 404. En la migración a Azure el cambio es masivo: hay que reemplazar "API_BASE = '
    "'https://asisa-pc.vercel.app'\" por el nuevo dominio en cada bloque (.js) y en api/markup.js."
)

doc.add_heading('3.6 Scripts del frontend EDS (carpeta scripts/)', level=2)
add_table(
    ['Fichero', 'Función'],
    [
        ('scripts/aem.js', 'Boilerplate de Adobe EDS. Utilidades core: loadHeader/Footer, '
                            'decorateButtons/Icons/Sections/Blocks, loadSection, loadCSS, '
                            'loadFragment. No tocar salvo upgrades.'),
        ('scripts/scripts.js', 'Orquesta la carga en 3 fases: loadEager (primera sección + LCP), '
                                'loadLazy (header, footer, lazy-styles), loadDelayed (importa '
                                'delayed.js a los 3s).'),
        ('scripts/delayed.js', 'Hueco para lógica no crítica (analytics, RUM). Actualmente vacío.'),
        ('scripts/editor-support.js', 'Sólo activo dentro del Universal Editor de AEM. Escucha '
                                       'eventos aue:content-* y re-decora bloques sin recargar.'),
        ('scripts/editor-support-rte.js', 'Agrupa nodos rich-text consecutivos para edición '
                                           'cómoda en el editor.'),
        ('scripts/dompurify.min.js', 'Librería externa de sanitización XSS, usada por '
                                      'editor-support.js.'),
    ],
    widths=[Inches(2), Inches(4.5)],
)

doc.add_heading('3.7 Estilos y assets', level=2)
add_table(
    ['Path', 'Función'],
    [
        ('styles/styles.css', 'Stub. Estilos vienen de las clientlibs de ASISA.'),
        ('styles/fonts.css', 'Vacío. Fuentes vía clientlibs.'),
        ('styles/lazy-styles.css', 'Placeholder para CSS no crítico (post-LCP).'),
        ('fonts/roboto-*.woff2', 'Roboto Regular/Medium/Bold/Condensed Bold como fallback local.'),
        ('icons/search.svg', 'Icono de búsqueda local; el resto desde clientlib-iconslib.'),
        ('favicon.ico', 'Favicon del site.'),
    ],
    widths=[Inches(2), Inches(4.5)],
)

# =============================================================================
# 4. Datos
# =============================================================================
doc.add_heading('4. Datos cacheados (carpeta data/)', level=1)
doc.add_paragraph(
    'Las funciones serverless NO consultan la API de ASISA en cada request. En su lugar leen JSON '
    'pre-generados que viven en el repo de GitHub. Eso da latencia mínima y desacopla la web del '
    'backend de ASISA.'
)
add_table(
    ['Path', 'Contenido', 'Volumen aprox.'],
    [
        ('data/provincias.json', 'Catálogo maestro de provincias (name, displayName, slug, code).', '52'),
        ('data/cuadro-medico/especialidades.json', 'Master list de especialidades.', '181'),
        ('data/cuadro-medico/doctores-index.json', 'slug-doctor → [ubicaciones].', '~20.500'),
        ('data/cuadro-medico/centros-index.json', 'slug-centro → [ubicaciones].', '~6.500'),
        ('data/cuadro-medico/especialidades/<spec>.json', 'Provincias donde está disponible.', '~183'),
        ('data/cuadro-medico/provincias/<prov>.json', 'Especialidades disponibles en la provincia.', '50'),
        ('data/providers/<prov>/<spec>.json', 'Lista cruda por (provincia, especialidad).', '~3.700'),
        ('data/provider-details/<locCode>.json', 'Detalle por código de localización.', '~33.000'),
    ],
    widths=[Inches(2.5), Inches(2.8), Inches(1.2)],
)
doc.add_paragraph(
    'En Vercel, estos JSON viajan dentro del bundle de cada función serverless. En Azure hay dos '
    'opciones: (a) seguir empaquetándolos con cada Function (simple, pero el bundle es grande), o '
    '(b) moverlos a Azure Blob Storage y que las funciones los lean por HTTP. La opción (b) es '
    'preferible para deployments más rápidos.'
)

doc.add_heading('4.1 Schemas detallados de cada fichero', level=2)

doc.add_heading('4.1.1 data/provincias.json', level=3)
doc.add_paragraph('1 fichero. Array de 52 entradas (España completa + Ceuta y Melilla). '
                  'Mantenido a mano; rara vez cambia.')
add_code_block(
    '[\n'
    '  {\n'
    '    "name": "MADRID",          // mayúsculas, como viene de ASISA\n'
    '    "displayName": "Madrid",   // título, para pintar en UI\n'
    '    "slug": "madrid",          // kebab-case, usado en URLs\n'
    '    "provinceCode": "28"       // código INE\n'
    '  },\n'
    '  …\n'
    ']'
)
doc.add_paragraph('Generador: manual. Consumido por: TODOS los scripts .mjs (filtrar por '
                  'provinceCode o slug) + api/providers.js + api/provincias.js.')

doc.add_heading('4.1.2 data/cuadro-medico/especialidades.json', level=3)
doc.add_paragraph('1 fichero. Master list de 181 especialidades.')
add_code_block(
    '[\n'
    '  {\n'
    '    "slug": "cardiologia",\n'
    '    "name": "Cardiología",                   // pretty\n'
    '    "nameApi": "CARDIOLOGÍA",                // como ASISA lo devuelve\n'
    '    "professionalPlural": "Cardiólogos",\n'
    '    "professionalPluralLower": "cardiólogos",\n'
    '    "kind": "specialty" | "service" | "technique",\n'
    '    "specialityCode": 9                      // numérico ASISA (solo si specialty)\n'
    '  },\n'
    '  …\n'
    ']'
)
doc.add_paragraph(
    'Generador: manual (extraído originalmente del backend ASISA). Consumido por: '
    'api/especialidades.js + generate-cuadro-medico-specs.mjs (validación de slugs). El campo '
    'kind se usa para filtrar: el bloque cuadro-medico-otras-especialidades excluye las "service" '
    'cuando muestra el top 15 nacional.'
)

doc.add_heading('4.1.3 data/cuadro-medico/doctores-index.json', level=3)
doc.add_paragraph('1 fichero. Object con ~20.500 keys. Es el ÍNDICE MAESTRO de profesionales. '
                  'Cada key tiene formato "<slug-nombre>-<collegiateCode|providerCode>".')
add_code_block(
    '{\n'
    '  "garcia-balda-ainhoa-152854071": {\n'
    '    "collegiateCode": 152854071,    // único por profesional (puede ser null)\n'
    '    "name": "GARCIA BALDA, AINHOA",\n'
    '    "locations": [\n'
    '      {\n'
    '        "providerCode": 7654321,\n'
    '        "providerLocalicationCode": 1234567,\n'
    '        "specSlug": "alergologia",\n'
    '        "provinceSlug": "a-coruna"\n'
    '      }\n'
    '      // … una ubicación por cada {centro, especialidad}\n'
    '    ]\n'
    '  },\n'
    '  …\n'
    '}'
)
doc.add_paragraph('Generador: generate-cuadro-medico-specs.mjs (agrupa entradas por '
                  'collegiateCode). Consumido por: api/doctor.js + create-aem-pages.mjs (un page '
                  '/cuadro-medico/d/<key> por cada key) + refresh-eds-pages.mjs --doctores.')

doc.add_heading('4.1.4 data/cuadro-medico/centros-index.json', level=3)
doc.add_paragraph('1 fichero. Object con ~6.500 keys. Cada key es el slug del centro.')
add_code_block(
    '{\n'
    '  "hm-rosaleda": {\n'
    '    "providerLocalicationCode": 12345,\n'
    '    "name": "HM ROSALEDA",\n'
    '    "provinceSlug": "a-coruna"\n'
    '  },\n'
    '  …\n'
    '}'
)
doc.add_paragraph('Generador: generate-cuadro-medico-specs.mjs. Consumido por: api/centro.js + '
                  'create-aem-pages.mjs + refresh-eds-pages.mjs --centros.')

doc.add_heading('4.1.5 data/cuadro-medico/provincias/<slug>.json', level=3)
doc.add_paragraph('50 ficheros (uno por provincia que tiene catálogo). 1 por provincia. Indica '
                  'QUÉ especialidades hay disponibles ahí.')
add_code_block(
    '{\n'
    '  "slug": "madrid",\n'
    '  "displayName": "Madrid",\n'
    '  "provinceCode": "28",\n'
    '  "especialidades": [\n'
    '    "alergologia", "cardiologia", "dermatologia-medico-quirurgica-y-venereo",\n'
    '    …\n'
    '  ]\n'
    '}'
)
doc.add_paragraph('Generador: generate-cuadro-medico-specs.mjs. Consumido por: '
                  'api/sitemap-cuadro-medico.js, refresh-eds-pages.mjs (genera todas las '
                  'combinaciones /p/<prov>/pe/<spec>), create-aem-pages.mjs.')

doc.add_heading('4.1.6 data/cuadro-medico/especialidades/<slug>.json', level=3)
doc.add_paragraph('~183 ficheros. Uno por especialidad. Indica EN QUÉ provincias está '
                  'disponible y con cuántos profesionales.')
add_code_block(
    '{\n'
    '  "slug": "cardiologia",\n'
    '  "provincias": [\n'
    '    { "slug": "madrid", "displayName": "Madrid", "count": 381 },\n'
    '    { "slug": "barcelona", "displayName": "Barcelona", "count": 218 },\n'
    '    …  // ordenado por count desc\n'
    '  ]\n'
    '}'
)
doc.add_paragraph('Generador: generate-cuadro-medico-specs.mjs. Consumido por: '
                  'api/especialidades.js, cuadro-medico-otras-provincias (bloque), '
                  'api/sitemap-cuadro-medico.js (type=especialidades).')

doc.add_heading('4.1.7 data/providers/<prov>/<spec>.json (LISTA CRUDA)', level=3)
doc.add_paragraph(
    '~3.700 ficheros distribuidos en 52 subcarpetas (una por provincia). Cada fichero contiene '
    'el array crudo de todos los providers (médicos+centros) que ASISA devolvió para esa '
    'combinación. Es el dataset MÁS PESADO y la fuente de TODO el resto de datos.'
)
add_code_block(
    '[\n'
    '  {\n'
    '    "providerCode": 7654321,                  // ID global del provider\n'
    '    "providerLocalicationCode": 1234567,      // ID único por ubicación\n'
    '    "providerName": "DR. PEREZ GARCIA",       // O "HOSPITAL HM …"\n'
    '    "networkCode": 1,                         // Red Salud\n'
    '    "doctorType": 1,                          // 1 = profesional, 0 = centro\n'
    '    "providerType": Number,                   // 3=HOSPITAL, 4=CENTRO,\n'
    '                                               //  8=LABORATORIO, etc.\n'
    '    "collegiateCode": 152854071,              // solo si doctorType=1\n'
    '    "gender": "M" | "F" | "",\n'
    '    "businessGroup": Boolean,\n'
    '    "specialityInfo": {\n'
    '      "specialityCode": 9,\n'
    '      "specialityDescription": "CARDIOLOGÍA",\n'
    '      "subSpecialityCode": Number,\n'
    '      "subSpecialityDescription": null | String\n'
    '    },\n'
    '    "contact": {\n'
    '      "phone": "915732022",\n'
    '      "mobilePhone": "",\n'
    '      "email": "",\n'
    '      "documentNumber": "12345678X"           // DNI del titular\n'
    '    },\n'
    '    "address": {\n'
    '      "addressType": "CL",                    // tipo de vía\n'
    '      "addressDescription": "DOCTOR ESQUERDO",\n'
    '      "addressNumber": 83,\n'
    '      "cityDescription": "MADRID",\n'
    '      "postalCode": "28028",\n'
    '      "provinceCode": 28,\n'
    '      "latitude": 40.43,\n'
    '      "longitude": -3.66\n'
    '    },\n'
    '    "parentDescription": "Hospital HM Universitario …",\n'
    '    "parentCode": Number,                     // → centro padre\n'
    '    "onlineAppointment": Boolean,\n'
    '    "videoConsultation": Boolean,\n'
    '    "electronicPrescription": Boolean,\n'
    '    "languages": [{ "code": "es" }, { "code": "en" }],\n'
    '    "tuotempo": {                             // datos del booking partner\n'
    '      "centerCode": Number,\n'
    '      "providerCode": Number,\n'
    '      "onlineAppointment": Boolean,\n'
    '      "presentialAppointment": Boolean,\n'
    '      "videoAppointment": Boolean,\n'
    '      "phoneAppointment": Boolean,\n'
    '      "asisaLiveAppointment": Boolean\n'
    '    }\n'
    '  },\n'
    '  …\n'
    ']'
)
doc.add_paragraph('Generador: generate-providers-data.mjs (descarga de ASISA). Consumido por: '
                  'TODO — api/providers.js (lectura directa), api/centro.js (scan), '
                  'generate-cuadro-medico-specs.mjs (construye los índices), '
                  'generate-provider-details.mjs (extrae locCodes únicos).')

doc.add_heading('4.1.8 data/provider-details/<locCode>.json', level=3)
doc.add_paragraph('~33.000 ficheros (uno por providerLocalicationCode único). Cada fichero '
                  'contiene un ARRAY (no objeto) con 1+ entries que comparten ubicación pero '
                  'pueden diferir en especialidad. Si la API ASISA falló, el fichero contiene '
                  '"null".')
add_code_block(
    '// data/provider-details/1234567.json\n'
    '[\n'
    '  {\n'
    '    "providerCode": 7654321,\n'
    '    "providerLocalicationCode": 1234567,\n'
    '    "providerName": "DR. PEREZ GARCIA",\n'
    '    "specialityInfo": { … },\n'
    '    "contact": { phone, email, documentNumber, … },\n'
    '    "address": { … },\n'
    '    // Campos ADICIONALES respecto a la lista cruda:\n'
    '    "observations": "Consulta privada en planta 3 …",\n'
    '    "appointments": Boolean,\n'
    '    // Horarios, servicios extra, etc. según el endpoint\n'
    '  }\n'
    '  // … más entries si el provider trabaja varias specs en esa ubicación\n'
    ']'
)
doc.add_paragraph('Generador: generate-provider-details.mjs (~33k fetches a /providers/details). '
                  'Consumido por: api/providers-detail.js + api/doctor.js (enriquece datos del '
                  'representative location).')

doc.add_heading('4.1.9 Flujo de datos entre todos los ficheros', level=3)
add_code_block(
    'API ASISA (ursaepre.asisa.es)\n'
    '  │\n'
    '  ├─ /searchPortal/providers ── generate-providers-data.mjs ──┐\n'
    '  │                                                            ▼\n'
    '  │                          data/providers/<prov>/<spec>.json (raw)\n'
    '  │                                                            │\n'
    '  ├─ /providers/details ── generate-provider-details.mjs ──┐  │\n'
    '  │                                                        ▼  │\n'
    '  │              data/provider-details/<locCode>.json         │\n'
    '  │                                                            │\n'
    '                                                                ▼\n'
    '                                  generate-cuadro-medico-specs.mjs\n'
    '                                                                │\n'
    '                                                                ▼\n'
    '                  data/cuadro-medico/\n'
    '                    ├─ provincias/<slug>.json\n'
    '                    ├─ especialidades/<slug>.json\n'
    '                    ├─ doctores-index.json\n'
    '                    └─ centros-index.json\n'
    '                                                                │\n'
    '                                                                ▼\n'
    '                  api/* (serverless) los leen en runtime'
)

# =============================================================================
# 5. Sitemaps
# =============================================================================
doc.add_heading('5. Sitemaps (configuración EDS-native)', level=1)
doc.add_paragraph(
    'La estrategia actual es 100% nativa de EDS: cada tipo de URL tiene su propio query-index y '
    'su propio sitemap. EDS los autogenera al servir las rutas /sitemap-*.xml a partir de la '
    'configuración en helix-query.yaml y helix-sitemap.yaml. Las URLs internas del XML apuntan a '
    'https://www.asisa.es gracias a la directiva origin.'
)
doc.add_paragraph('Sitemaps generados:')
add_table(
    ['URL', 'Origen', 'Contenido'],
    [
        ('/sitemap.xml',
         'pages-static index (estáticas)',
         'Páginas autoradas en AEM, no las dinámicas del cuadro médico.'),
        ('/sitemap-cuadro-medico-provincias.xml', 'cuadro-medico-provincias index', '52 URLs /p/<slug>'),
        ('/sitemap-cuadro-medico-provincia-specs.xml',
         'cuadro-medico-provincia-specs index', '~3.250 URLs /p/<prov>/pe/<spec>'),
        ('/sitemap-cuadro-medico-doctores.xml',
         'cuadro-medico-doctores index', '~20.500 URLs /d/<key>'),
        ('/sitemap-cuadro-medico-centros.xml',
         'cuadro-medico-centros index', '~6.500 URLs /c/<key>'),
        ('/sitemap-cuadro-medico-especialidades.xml',
         'cuadro-medico-especialidades index', '~181 URLs /e/<slug>'),
    ],
    widths=[Inches(2.5), Inches(2.2), Inches(1.8)],
)
doc.add_paragraph('Importante: las respuestas vienen comprimidas con Brotli desde la CDN. Curl '
                  'sin --compressed devuelve bytes binarios; usar:')
add_code_block('curl --compressed https://main--asisa-pc--asisa-softtek.aem.live/sitemap-cuadro-medico-doctores.xml')
doc.add_paragraph(
    'Las cachés api/sitemap.js y api/sitemap-cuadro-medico.js en Vercel se conservan como vía '
    'paralela útil para diagnóstico (accesibles directamente en https://<overlay-host>/sitemap.xml '
    'y similares), pero EDS no las consume. En la migración a Azure pueden eliminarse o portarse '
    'según se prefiera.'
)

doc.add_heading('5.1 Repoblar índices tras cambios', level=2)
doc.add_paragraph(
    'Cuando se cambia helix-query.yaml, los índices existentes NO se reescriben automáticamente; '
    'solo se actualizan a medida que se previsualiza cada path. Para forzar la repoblación masiva:'
)
add_code_block('node refresh-eds-pages.mjs --reindex')
doc.add_paragraph(
    'Este flag recorre las ~30.000 URLs en data/ y dispara POST /index/{path} contra '
    'admin.hlx.page con concurrencia 10. Tarda en torno a 40 minutos para el catálogo completo. '
    'No necesita token (anónimo gracias a requireAuth: auto).'
)

# =============================================================================
# 6. Scripts de mantenimiento
# =============================================================================
doc.add_heading('6. Scripts de mantenimiento (raíz del repo)', level=1)
doc.add_paragraph(
    'Toda la cadena de regeneración de datos y publicación está en scripts node ESM en la raíz:'
)
add_table(
    ['Script', 'Función'],
    [
        ('generate-providers-data.mjs',
         'Descarga el catálogo desde la API de ASISA (/searchPortal) por provincia y especialidad. '
         'Concurrencia 10. Flags: FORCE=true, PROVINCE_CODE=28. Salida: data/providers/.'),
        ('generate-provider-details.mjs',
         'Pre-descarga el detalle de cada (locCode, docNum) único desde /providers/details. '
         'Concurrencia 25 (bajar si 429). Salida: data/provider-details/.'),
        ('generate-cuadro-medico-specs.mjs',
         'Agrega lo anterior y construye los índices doctores-index.json, centros-index.json, '
         'especialidades/*.json y provincias/*.json.'),
        ('create-aem-pages.mjs',
         'Crea las páginas plantilla en AEM Author (copy + publish) y refresca EDS. Requiere '
         'AEM_TOKEN. Flags: --provincias, --especialidades, --doctores, --centros, --all.'),
        ('refresh-eds-pages.mjs',
         'Refresca preview+live (y opcionalmente reindex) en EDS sin tocar AEM. Flags: --code, '
         '--provincias, --specs, --doctores, --centros, --especialidades, --sitemaps, --reindex, '
         '--province=madrid. Requiere HLX_ADMIN_API_TOKEN.'),
    ],
    widths=[Inches(2.5), Inches(4)],
)
doc.add_paragraph('Flujo encadenado típico para una actualización completa:')
add_code_block(
    '# 1. Descargar catálogo desde la API de ASISA\n'
    'node generate-providers-data.mjs\n\n'
    '# 2. Descargar detalles por localización\n'
    'node generate-provider-details.mjs\n\n'
    '# 3. Construir índices\n'
    'node generate-cuadro-medico-specs.mjs\n\n'
    '# 4. (Solo la primera vez) crear páginas plantilla en AEM — necesita AEM_TOKEN\n'
    'AEM_TOKEN=… node create-aem-pages.mjs --all\n\n'
    '# 5. Refrescar todo en EDS (anónimo, no necesita token)\n'
    'node refresh-eds-pages.mjs\n\n'
    '# 6. Repoblar índices si se cambió helix-query.yaml\n'
    'node refresh-eds-pages.mjs --reindex'
)

doc.add_heading('6.1 Detalle exhaustivo de cada script', level=2)

# generate-providers-data.mjs
doc.add_heading('Script 1 · generate-providers-data.mjs (~190 líneas)', level=3)
doc.add_paragraph(
    'Propósito: descarga el catálogo completo de profesionales y centros de ASISA, lo dedupla y '
    'lo guarda en data/providers/<provinciaSlug>/<especialidadSlug>.json. Es el cimiento de '
    'toda la pirámide de datos.'
)
doc.add_paragraph('Inputs:')
add_bullet('Env: FORCE=true (sobrescribe caché existente); PROVINCE_CODE=28 (solo una provincia).')
add_bullet('Ficheros: data/provincias.json (lista provincias y sus códigos).')
doc.add_paragraph('Outputs: data/providers/<prov>/<spec>.json — arrays planos de providers.')
doc.add_paragraph('APIs externas a ASISA:')
add_code_block(
    'Base: https://ursaepre.asisa.es/ASISA/middlewasisa/public/v1/api/searchPortal\n\n'
    '1) GET /autocomplete/specialities\n'
    '   ?specialityDescription=&networkCode=1&provinceCode=<code>&maxResultsNumber=200\n'
    '   → lista de especialidades disponibles en esa provincia\n\n'
    '2) GET /providers\n'
    '   ?networkCode=1&provinceCode=<code>\n'
    '   &specialityDescription=<DESC>&specialityType=<N>\n'
    '   &pageNumber=<P>      // 100 results por página\n'
    '   → listado paginado de providers\n\n'
    'Headers (todas):\n'
    '  Ocp-Apim-Subscription-Key: 0908b85b9d0e4a75b2eb33048bd9fe01\n'
    '  Api-Version: 1\n\n'
    'Timeout: 150 s'
)
doc.add_paragraph('Flujo:')
add_numbered('Lee data/provincias.json y filtra por PROVINCE_CODE si se ha pasado.')
add_numbered('Para cada provincia (concurrencia 10), llama /autocomplete/specialities para '
             'obtener qué especialidades tiene.')
add_numbered('Para cada (provincia, especialidad), llama /providers paginando de 100 en 100 hasta '
             'consumir totalCount.')
add_numbered('Deduplica usando un Set por providerCode (si existe) o fallback (name, address, '
             'city).')
add_numbered('Si el fichero destino existe y no hay FORCE → skip (lee caché del repo).')
add_numbered('Escribe data/providers/<prov>/<spec>.json.')
doc.add_paragraph('Concurrencia: 10. Sin retry automático para 429 (timeout 150 s ya es generoso). '
                  'Errores tolerados: status != 200, NETWORK, TIMEOUT — log y continúa.')
doc.add_paragraph('Lanzador: manual (CLI) o GitHub Action workflow_dispatch (con cron diario '
                  'disponible pero comentado).')

# generate-provider-details.mjs
doc.add_heading('Script 2 · generate-provider-details.mjs (~158 líneas)', level=3)
doc.add_paragraph('Propósito: por cada (providerLocalicationCode, documentNumber) único en los '
                  'JSON de providers, llama /providers/details para obtener datos ampliados '
                  '(observaciones, horarios, etc.) y los guarda en data/provider-details/.')
doc.add_paragraph('Inputs: env FORCE; lee data/providers/**/*.json para extraer los pares únicos.')
doc.add_paragraph('Outputs: data/provider-details/<providerLocalicationCode>.json — array con 1+ '
                  'entries (o null si error).')
doc.add_paragraph('APIs externas:')
add_code_block(
    'GET /searchPortal/providers/details\n'
    '  ?networkCode=1&providerLocalicationCode=<id>&documentNumber=<dni>\n'
    'Headers iguales que generate-providers-data.\n'
    'Timeout: 150 s.'
)
doc.add_paragraph('Concurrencia: 25 (bajar a 10 si la API empieza a devolver 429). Si la '
                  'request falla, escribe "null" en el fichero (no propaga el error).')
doc.add_paragraph('Lanzador: manual o workflow_dispatch (timeout configurado a 360 min, '
                  '~33k peticiones).')

# generate-cuadro-medico-specs.mjs
doc.add_heading('Script 3 · generate-cuadro-medico-specs.mjs (~236 líneas)', level=3)
doc.add_paragraph(
    'Propósito: NO toca la red. Lee TODOS los data/providers/<prov>/<spec>.json y produce los '
    '4 ficheros de índice + los 2 directorios de detalle agregado.'
)
doc.add_paragraph('Inputs: env PROVINCE_SLUG=madrid (solo una); lee provincias.json, '
                  'especialidades.json, data/providers/**/*.json.')
doc.add_paragraph('Outputs:')
add_bullet('data/cuadro-medico/provincias/<slug>.json (52)')
add_bullet('data/cuadro-medico/especialidades/<slug>.json (~183)')
add_bullet('data/cuadro-medico/doctores-index.json (~20.500 keys)')
add_bullet('data/cuadro-medico/centros-index.json (~6.500 keys)')
doc.add_paragraph('Flujo:')
add_numbered('Itera cada (provincia, especialidad) → lee el JSON crudo.')
add_numbered('Para cada provider: si doctorType=1, key="<slug-nombre>-<collegiateCode|providerCode>" '
             '→ doctoresIndex. Si no, key="<slug-nombre>" → centrosIndex.')
add_numbered('Mantiene un mapa provSlug → Set(specSlugs) y otro specSlug → Map(provSlug → count).')
add_numbered('Al final escribe los 4 ficheros / 2 carpetas con los agregados.')
doc.add_paragraph('Errores tolerados: ficheros vacíos o corruptos → skip silencioso. Fatales: '
                  'sin provincias.json o sin especialidades.json → exit(1).')
doc.add_paragraph('Lanzador: manual tras los dos generate-*.mjs anteriores.')

# create-aem-pages.mjs
doc.add_heading('Script 4 · create-aem-pages.mjs (~183 líneas)', level=3)
doc.add_paragraph(
    'Propósito: crea en AEM Author la página AEM correspondiente a cada URL dinámica (copy desde '
    'plantilla), la activa (publish), y refresca preview+live en EDS. Es el único script que '
    'toca AEM directamente.'
)
doc.add_paragraph('Inputs:')
add_bullet('Env: AEM_TOKEN="login:eyJ..." (REQUERIDO; cookie login-token del Author).')
add_bullet('Flags: --provincias, --especialidades, --doctores, --centros, --all.')
add_bullet('Lee: provincias.json + provincias/<slug>.json + especialidades/ + doctores-index '
           '+ centros-index según los flags.')
doc.add_paragraph('APIs externas:')
add_code_block(
    'AEM Author (https://author-p133185-e1320482.adobeaemcloud.com)\n'
    '  GET  {path}.infinity.json          → comprueba si la página existe\n'
    '  POST /bin/wcmcommand               → cmd=copyPage&srcPath=…&destParentPath=…&pageName=…\n'
    '                                      &shallow=false&replaceExistingPages=false\n'
    '  POST /bin/replicate.json           → cmd=activate&path=…\n'
    '  Headers: Cookie: login-token=<AEM_TOKEN>\n'
    '           Content-Type: application/x-www-form-urlencoded\n\n'
    'EDS admin (https://admin.hlx.page)\n'
    '  POST /preview/asisa-softtek/asisa-pc/main<edsPath>\n'
    '  POST /live/asisa-softtek/asisa-pc/main<edsPath>'
)
doc.add_paragraph('Plantillas AEM de las que copia:')
add_bullet('/content/site-pc/cuadro-medico/provincia → /cuadro-medico/p/<slug>')
add_bullet('/content/site-pc/cuadro-medico/provincia → /cuadro-medico/p/<prov>/pe/<spec>')
add_bullet('/content/site-pc/cuadro-medico/doctor → /cuadro-medico/d/<key>')
add_bullet('/content/site-pc/cuadro-medico/centro → /cuadro-medico/c/<key>')
add_bullet('/content/site-pc/cuadro-medico/especialidad → /cuadro-medico/e/<slug>')
doc.add_paragraph('Concurrencia 5 con delay 200 ms entre batches (concurrencia más alta puede '
                  'generar locks OAK en AEM). Errores tolerados: FAIL copy → log y continúa. '
                  'Fatal: sin AEM_TOKEN → exit(1).')
doc.add_paragraph('Importante para la migración: este script es independiente de Vercel. '
                  'Funciona igual sea cual sea el overlay. No requiere cambios.')

# refresh-eds-pages.mjs
doc.add_heading('Script 5 · refresh-eds-pages.mjs (~254 líneas)', level=3)
doc.add_paragraph(
    'Propósito: refresca preview/live y opcionalmente repuebla los query-indexes en EDS, sin '
    'tocar AEM. Es el "Swiss army knife" del día a día.'
)
doc.add_paragraph('Inputs:')
add_bullet('Env: HLX_ADMIN_API_TOKEN (opcional, gracias a requireAuth: auto).')
add_bullet('Flags: --code, --provincias, --specs, --doctores, --centros, --especialidades, '
           '--sitemaps, --reindex, --province=madrid. Sin flags → MODO FULL (todo).')
add_bullet('Lee data/ para construir las listas de paths.')
doc.add_paragraph('APIs externas:')
add_code_block(
    'POST https://admin.hlx.page/code/asisa-softtek/asisa-pc/main\n'
    'POST https://admin.hlx.page/preview/asisa-softtek/asisa-pc/main<path>\n'
    'POST https://admin.hlx.page/live/asisa-softtek/asisa-pc/main<path>\n'
    'POST https://admin.hlx.page/index/asisa-softtek/asisa-pc/main<path>\n'
    'Headers: x-auth-token: <HLX_ADMIN_API_TOKEN>     // opcional con requireAuth: auto'
)
doc.add_paragraph('Flujo:')
add_numbered('Refresca code bus (siempre que se pase code o esté en modo full).')
add_numbered('Para cada tipo seleccionado: lee data/, construye lista de paths, '
             'POST /preview seguido de POST /live con delay 100 ms entre llamadas.')
add_numbered('Si --reindex: POST /index para cada path. Si la URL no existe en EDS '
             '(404), se loga y sigue.')
add_numbered('Si --sitemaps: dispara preview+live de los 6 paths de sitemap.')
doc.add_paragraph('Concurrencia: 10. Retry: 3 intentos con backoff 500 ms × intento para '
                  'errores no fatales. 401 y 404 → no se reintenta.')
doc.add_paragraph('Tiempos de referencia con el catálogo actual (concurrencia 10):')
add_bullet('--code: < 5 s.')
add_bullet('--provincias: < 1 min.')
add_bullet('--specs: 3-5 min.')
add_bullet('--doctores: 35-45 min.')
add_bullet('--centros: 10-15 min.')
add_bullet('--especialidades: < 1 min.')
add_bullet('--sitemaps: < 1 min.')
add_bullet('--reindex (todo): ~40 min.')

doc.add_heading('6.2 Comandos sueltos clave', level=2)
doc.add_paragraph(
    'Ninguno de estos comandos necesita token (requireAuth: auto en el access config).'
)
doc.add_paragraph('Publicar una URL concreta:')
add_code_block(
    '# Preview (lo procesa pero no es público)\n'
    'curl -X POST https://admin.hlx.page/preview/asisa-softtek/asisa-pc/main/cuadro-medico/p/madrid\n\n'
    '# Live (lo hace público, requiere preview previo)\n'
    'curl -X POST https://admin.hlx.page/live/asisa-softtek/asisa-pc/main/cuadro-medico/p/madrid'
)
doc.add_paragraph('Consultar el estado de una URL en EDS:')
add_code_block('curl https://admin.hlx.page/status/asisa-softtek/asisa-pc/main/cuadro-medico/p/madrid')
doc.add_paragraph('Refrescar el code bus (JS, CSS, head.html, YAMLs) tras un push a GitHub:')
add_code_block(
    '# Refresca todo el code bus\n'
    'curl -X POST https://admin.hlx.page/code/asisa-softtek/asisa-pc/main\n\n'
    '# Solo un fichero\n'
    'curl -X POST https://admin.hlx.page/code/asisa-softtek/asisa-pc/main/head.html'
)
doc.add_paragraph(
    'Si el GitHub App "AEM Code Sync" está instalado en el repo, el code bus se '
    'refresca solo en cada push. El curl manual es para forzarlo.'
)

# =============================================================================
# 7. GitHub Actions
# =============================================================================
doc.add_heading('7. GitHub Actions', level=1)
add_table(
    ['Workflow', 'Trigger', 'Función'],
    [
        ('.github/workflows/main.yaml',
         'Push a cualquier rama',
         'CI básico: npm ci + npm run lint (ESLint + Stylelint).'),
        ('.github/workflows/generate-providers-data.yml',
         'workflow_dispatch (manual, input force)',
         'Ejecuta generate-providers-data.mjs y commitea los cambios en data/. Cron diario '
         'comentado.'),
        ('.github/workflows/generate-provider-details.yml',
         'workflow_dispatch',
         'Ejecuta generate-provider-details.mjs. Timeout 360 min. Cron diario comentado.'),
    ],
    widths=[Inches(2.5), Inches(2), Inches(2)],
)

# =============================================================================
# 8. Migración a Azure — el capítulo clave
# =============================================================================
doc.add_page_break()
doc.add_heading('8. Migración de Vercel a Azure', level=1)
doc.add_paragraph(
    'Este capítulo es la guía operativa de la migración. La filosofía es: cambiar SOLO la capa '
    'Vercel; AEM Author, Edge Delivery Services y GitHub se quedan exactamente como están.'
)

doc.add_heading('8.1 Qué hace Vercel hoy (inventario completo)', level=2)
add_table(
    ['Capacidad', 'Endpoint actual', 'Implementación'],
    [
        ('Overlay BYOM (HTML plantillas)', 'https://asisa-pc.vercel.app/markup/*',
         'api/markup.js como Vercel Function'),
        ('API datos: listados', 'https://asisa-pc.vercel.app/api/providers',
         'api/providers.js'),
        ('API datos: ficha doctor', 'https://asisa-pc.vercel.app/api/doctor',
         'api/doctor.js'),
        ('API datos: ficha centro', 'https://asisa-pc.vercel.app/api/centro',
         'api/centro.js'),
        ('API datos: provincias', 'https://asisa-pc.vercel.app/api/provincias',
         'api/provincias.js'),
        ('API datos: especialidades', 'https://asisa-pc.vercel.app/api/especialidades',
         'api/especialidades.js'),
        ('API datos: providers-detail', 'https://asisa-pc.vercel.app/api/providers-detail',
         'api/providers-detail.js'),
        ('Proxy autocomplete ASISA', 'https://asisa-pc.vercel.app/api/specialities',
         'api/specialities.js (fetch en vivo a ursaepre.asisa.es)'),
        ('Sync sitemap → EDS', 'https://asisa-pc.vercel.app/api/sync-aem',
         'api/sync-aem.js'),
        ('Proxy clientlibs CSS de ASISA', 'https://asisa-pc.vercel.app/etc.clientlibs/*',
         'rewrite en vercel.json hacia www.asisa.es'),
        ('Hosting de los JSON cacheados', '~3.700 + 33.000 ficheros en data/',
         'Empaquetados con cada function en cada deploy'),
        ('CORS automático en /api/*', 'Headers de respuesta',
         'Vercel inyecta CORS por defecto'),
        ('Auto-deploy desde GitHub', '',
         'Conector Vercel ↔ GitHub (actualmente DESCONECTADO; el deploy de las últimas semanas se '
         'ha hecho manualmente con `vercel deploy --prod --archive=tgz`)'),
    ],
    widths=[Inches(2.2), Inches(2.5), Inches(2)],
)

doc.add_heading('8.2 Equivalencias Vercel → Azure', level=2)
add_table(
    ['Capacidad Vercel', 'Azure recomendado', 'Notas'],
    [
        ('Vercel Functions (Node 22, ESM)',
         'Azure Functions (Node 22, ESM) con runtime v4',
         'Mismo modelo de funciones. Hay que adaptar el "request handler shape" — Azure '
         'Functions usa context.res en lugar de (req, res). Wrappers fáciles de añadir.'),
        ('Edge Functions / static',
         'Azure Front Door + Azure Functions Premium',
         'Front Door da CDN global y reglas de routing/rewrite equivalentes a vercel.json.'),
        ('Rewrites de vercel.json',
         'Reglas en Azure Front Door (rules engine) o azure-functions host.json',
         'Para /etc.clientlibs/* el proxy se hace con una Front Door Origin pointing to '
         'www.asisa.es y una regla path-based.'),
        ('Headers CORS',
         'Configurar en host.json de Azure Functions o vía Front Door',
         'Definir Access-Control-Allow-Origin: * para /api/*.'),
        ('Lectura de data/*.json',
         'Azure Blob Storage (preferido) o bundled with function',
         'Si los datos cambian a menudo (workflow GH Actions), Blob Storage es mejor: '
         'el deploy de funciones queda independiente de los datos. Usar @azure/storage-blob.'),
        ('Caché in-memory por instancia',
         'Sigue funcionando en Azure Functions',
         'Premium plan = instancias más "warm". Considerar Azure Cache for Redis si se quiere '
         'caché compartida entre instancias.'),
        ('Auto-deploy desde GitHub',
         'GitHub Actions con azure/functions-action',
         'Workflow estándar: en push a main, build + deploy a Function App. Habilitar deployment '
         'slots para staging.'),
        ('Dominio asisa-pc.vercel.app',
         'CNAME a asisa-pc.azurefd.net (o custom)',
         'El dominio del overlay puede cambiar — solo hay que actualizar la URL en la config '
         'sitewide de EDS (§2.1) y en los bloques (§3.5).'),
        ('Variables de entorno',
         'Azure Function App Configuration / Key Vault',
         'SYNC_SECRET, HLX_ADMIN_API_TOKEN, ASISA API keys.'),
    ],
    widths=[Inches(2), Inches(2), Inches(2.5)],
)

doc.add_heading('8.3 Plan de migración paso a paso', level=2)
add_numbered('Crear el Resource Group en Azure (rg-asisa-pc-prod, rg-asisa-pc-stage).')
add_numbered('Provisionar Azure Storage Account para los JSON de data/. Subir el catálogo '
             'actual y crear un container "asisa-pc-data" público read-only o privado con SAS.')
add_numbered('Provisionar Azure Function App (Node 22, Premium plan ideal para evitar cold '
             'starts en /markup/*).')
add_numbered('Portar el código de api/*.js: cambiar la firma a Azure Functions v4 '
             '(app.http(name, handler)), reemplazar readFileSync de data/ por fetch al Blob '
             'Storage (o seguir empaquetando si se prefiere simplicidad).')
add_numbered('Provisionar Azure Front Door con un Origin Group apuntando a la Function App. '
             'Añadir reglas: (a) /etc.clientlibs/* → origin www.asisa.es, (b) /sitemap.xml y '
             '/sitemap-cuadro-medico-*.xml → Function App (opcional, EDS ya los genera nativos), '
             '(c) /markup/* → Function App.')
add_numbered('Configurar dominio custom en Front Door (ej. asisa-pc.azurefd.net).')
add_numbered('Crear un workflow GitHub Actions que despliega a Azure Functions en cada push.')
add_numbered('Actualizar las URLs en los bloques (blocks/*/*.js — buscar "asisa-pc.vercel.app" '
             'y reemplazar). Ojo: hay ~10 ficheros con la URL hardcodeada.')
add_numbered('Actualizar head.html: las tres líneas que cargan /etc.clientlibs/* deben '
             'apuntar al nuevo dominio Azure.')
add_numbered('Actualizar api/markup.js: el template HTML inyecta tres <link rel="stylesheet" '
             'href="https://asisa-pc.vercel.app/etc.clientlibs/..."> — cambiar al nuevo dominio.')
add_numbered('Hacer push a main (cambios de URL + nuevo workflow Azure).')
add_numbered('Ejecutar POST /code refresh para que EDS recoja head.html y los bloques nuevos.')
add_numbered('Modificar la sitewide config de EDS: cambiar content.overlay.url de '
             'https://asisa-pc.vercel.app/markup a https://asisa-pc.azurefd.net/markup '
             '(comando completo en §2.1). Requiere token config_admin.')
add_numbered('Ejecutar `node refresh-eds-pages.mjs --reindex` para que EDS reconstruya '
             'sus índices apuntando al overlay nuevo.')
add_numbered('Verificar: status de una URL de cuadro médico debe mostrar sourceLocation con '
             'el nuevo overlay; los bloques deben pintar datos correctamente; los 6 sitemaps '
             'deben responder 200 con XML válido.')
add_numbered('Una vez verificado, retirar el proyecto Vercel (o dejarlo apagado un tiempo '
             'como backup).')

doc.add_heading('8.4 Puntos delicados de la migración', level=2)
add_bullet('Cold starts: Azure Functions en Consumption Plan puede tener cold starts >2s, '
           'inaceptable para el overlay BYOM (EDS hace request y el HTML aparece en LCP). '
           'Usar Premium Plan o Always-On.')
add_bullet('Imports entre funciones: api/markup.js importa de api/sitemap.js y api/sitemap-cuadro-medico.js. '
           'En Vercel cada función se bundlea independiente; en Azure Functions v4 hay que '
           'configurar el packaging para compartir código vía shared/.')
add_bullet('Tamaño del deploy: con 33.000 ficheros JSON en data/ el bundle de cada función es '
           'enorme. Si NO se usa Blob Storage, hay que aumentar los límites de deploy en Azure '
           '(default ~250MB).')
add_bullet('CORS: las APIs necesitan Access-Control-Allow-Origin: * porque los bloques en aem.live '
           'hacen fetch cross-origin. Configurar explícitamente en host.json o vía Front Door rules.')
add_bullet('Permisos EDS: hace falta el token config_admin (cuenta técnica de Adobe) para '
           'cambiar la sitewide config. jorge.lorenzo@ext.softtek.com NO tiene ese rol. Localizar '
           'al técnico de Adobe (proyecto en https://developer.adobe.com/console) antes de '
           'empezar la migración.')
add_bullet('Rollback: mientras Azure no esté validado, mantener Vercel funcional. Para volver, '
           'basta con un nuevo POST sitewide config apuntando de nuevo al overlay Vercel.')

# =============================================================================
# 9. Operación día a día
# =============================================================================
doc.add_page_break()
doc.add_heading('9. Operación día a día', level=1)

doc.add_heading('9.1 Tras cambiar JS o CSS de un bloque', level=2)
add_numbered('Commit + push a main.')
add_numbered('Si AEM Code Sync está activo en el repo, EDS se entera solo. Si no: '
             'POST /code refresh.')
add_numbered('NO hace falta re-publicar páginas; los bloques se ejecutan en cliente.')

doc.add_heading('9.2 Tras cambiar head.html', level=2)
add_numbered('Commit + push.')
add_numbered('POST /code refresh.')
add_numbered('POST /preview y /live para CADA URL afectada — head.html se inlinea en el HTML '
             'procesado, no se sirve dinámicamente. Para masivo: '
             '`node refresh-eds-pages.mjs` (sin flags = todo).')

doc.add_heading('9.3 Tras cambiar la plantilla en api/markup.js', level=2)
add_numbered('Commit + push.')
add_numbered('Deploy a Vercel (o Azure tras migración).')
add_numbered('`node refresh-eds-pages.mjs` para re-previewar todas las URLs dinámicas — EDS '
             're-fetcha del overlay.')

doc.add_heading('9.4 Tras actualizar datos (data/*.json)', level=2)
add_numbered('Commit + push (los workflows GH Actions ya lo hacen).')
add_numbered('Vercel/Azure recoge los datos en el siguiente deploy. Si Blob Storage, los datos '
             'son visibles inmediatamente sin redeploy.')
add_numbered('Los bloques fetchean datos en runtime, así que NO hace falta re-publicar páginas.')

doc.add_heading('9.5 Verificación rápida del estado del sistema', level=2)
add_code_block(
    '# Health check del overlay\n'
    'curl -sI https://asisa-pc.vercel.app/markup/cuadro-medico/p/madrid\n'
    '# → HTTP 200, x-source: template:provincia\n\n'
    '# Health check de una API\n'
    'curl -s "https://asisa-pc.vercel.app/api/provincias" | jq "length"\n'
    '# → 52\n\n'
    '# Status de una URL en EDS\n'
    'curl https://admin.hlx.page/status/asisa-softtek/asisa-pc/main/cuadro-medico/p/madrid\n'
    '# → preview/live status 200, sourceLocation apuntando al overlay\n\n'
    '# Sitemap de un tipo\n'
    'curl --compressed https://main--asisa-pc--asisa-softtek.aem.live/sitemap-cuadro-medico-provincias.xml\n'
    '# → XML con 52 <url>'
)

# =============================================================================
# 10. Troubleshooting
# =============================================================================
doc.add_heading('10. Troubleshooting', level=1)

add_table(
    ['Síntoma', 'Causa probable', 'Solución'],
    [
        ('Página dinámica devuelve 404 desde aem.live',
         'La URL no ha sido previsualizada/publicada en EDS',
         'POST /preview + POST /live para esa URL (o node refresh-eds-pages.mjs)'),
        ('Bloque no carga datos ("No se pudieron cargar")',
         'fetch relativo en lugar de absoluto',
         'Cambiar fetch("/api/...") por fetch("https://<overlay>/api/...")'),
        ('Página "en blanco" sin estilos',
         'Falta las clientlibs CSS de ASISA',
         'Comprobar que api/markup.js inyecta los <link> a /etc.clientlibs/* y que el rewrite '
         'apunta correctamente'),
        ('Cambio en api/markup.js no se ve',
         'EDS cachea el HTML procesado',
         'POST /preview + POST /live para cada URL'),
        ('POST /config devuelve 401',
         'Usuario no está en config_admin',
         'Usar el token de la cuenta técnica de Adobe'),
        ('/query-index.json devuelve 413',
         'Más de 6MB (límite de Lambda interno de EDS)',
         'Reducir vía exclude en helix-query.yaml o split en varios índices (ya hecho en este proyecto)'),
        ('Vercel no auto-despliega tras push',
         'El conector GitHub-Vercel está desconectado',
         'Lanzar manualmente `vercel deploy --prod --archive=tgz`'),
        ('Sitemap viene en bytes binarios',
         'Comprimido en Brotli',
         'Añadir --compressed al curl'),
        ('POST /preview/sitemap-*.xml devuelve 415',
         'El content-bus de EDS rechaza XML por el pipeline de preview',
         'No es bloqueante: EDS genera los sitemaps nativamente vía helix-sitemap.yaml sin '
         'necesidad de preview/live'),
    ],
    widths=[Inches(2.2), Inches(2), Inches(2.3)],
)

# =============================================================================
# 11. Apéndices
# =============================================================================
doc.add_heading('11. Apéndices', level=1)

doc.add_heading('11.1 URLs y dominios clave', level=2)
add_table(
    ['Entorno', 'URL'],
    [
        ('EDS Live (público)', 'https://main--asisa-pc--asisa-softtek.aem.live'),
        ('EDS Preview (autenticado)', 'https://main--asisa-pc--asisa-softtek.aem.page'),
        ('Vercel (overlay + APIs)', 'https://asisa-pc.vercel.app'),
        ('AEM Author', 'https://author-p133185-e1320482.adobeaemcloud.com'),
        ('Admin HLX API', 'https://admin.hlx.page'),
        ('GitHub repo', 'https://github.com/asisa-softtek/asisa-pc'),
        ('Producción final (target)', 'https://www.asisa.es'),
    ],
    widths=[Inches(2.5), Inches(4)],
)

doc.add_heading('11.2 Variables de entorno y secretos', level=2)
add_table(
    ['Variable', 'Uso', 'Dónde se configura'],
    [
        ('HLX_ADMIN_API_TOKEN', 'Token admin EDS para POST /preview, /live, /code, /index',
         'Local (export) o GitHub Secret'),
        ('AEM_TOKEN', 'Token de AEM Author para create-aem-pages.mjs',
         'Local (export)'),
        ('SYNC_SECRET', 'Shared secret para api/sync-aem.js',
         'Vercel env / Azure Function App settings'),
        ('FORCE', 'Flag para forzar regeneración en los generate-*.mjs',
         'CLI inline o input del workflow'),
        ('PROVINCE_CODE', 'Restringe el scrape a una provincia',
         'CLI inline'),
    ],
    widths=[Inches(2), Inches(2.5), Inches(2)],
)

doc.add_heading('11.3 Docs internas del proyecto', level=2)
add_bullet('docs/byom.md — explicación detallada del setup BYOM, los 3 POSTs originales, '
           'troubleshooting histórico y permisos.')
add_bullet('docs/estructura-proyecto.md — referencia fichero por fichero de todo el repo, con '
           'URLs en vivo donde inspeccionar cada pieza.')
add_bullet('docs/traspaso-conocimiento.docx — este documento.')

doc.add_heading('11.4 Glosario rápido', level=2)
add_table(
    ['Término', 'Significado'],
    [
        ('EDS', 'Edge Delivery Services (Adobe). CDN + capa de entrega para sites AEM.'),
        ('BYOM', 'Bring Your Own Markup. Patrón donde EDS pide HTML a un overlay externo en lugar '
                 'de a AEM Author.'),
        ('Overlay', 'Servidor que devuelve el HTML para URLs dinámicas. Hoy en Vercel, mañana en '
                    'Azure.'),
        ('Content bus', 'Almacenamiento interno de EDS para los HTML procesados (preview y live).'),
        ('Code bus', 'Almacenamiento interno de EDS para el código del repo (.js, .css, .yaml).'),
        ('Sidekick', 'Toolbar de edición integrada en el navegador para autores.'),
        ('Universal Editor', 'Editor visual de AEM Author donde se autora el contenido estático.'),
        ('Clientlib', 'Librería compilada de CSS/JS publicada por AEM. En este proyecto, las del '
                       'design system de ASISA viven en www.asisa.es/etc.clientlibs/wasisa/.'),
        ('Bloque', 'Componente EDS instanciado a partir de <div class="<nombre>">. Vive en '
                    'blocks/<nombre>/.'),
        ('Query index', 'JSON con los paths indexados de páginas (origen de los sitemaps).'),
    ],
    widths=[Inches(2), Inches(4.5)],
)

# =============================================================================
# Save
# =============================================================================
doc.save(OUTPUT)
print(f"Generado: {OUTPUT}")
print(f"Tamaño: {os.path.getsize(OUTPUT):,} bytes")
